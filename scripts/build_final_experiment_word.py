from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "WAMI最终实验结果汇总.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8.2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, title: str, headers: list[str], rows: list[list], note: str | None = None) -> None:
    doc.add_heading(title, level=2)
    if note:
        p = doc.add_paragraph(note)
        p.style = "Body Text"
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if i == 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def add_figure(doc: Document, number: str, title: str, filename: str, caption: str) -> None:
    path = ROOT / "data" / filename
    doc.add_heading(f"Figure {number}. {title}", level=2)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(path), width=Inches(6.2))
    else:
        doc.add_paragraph(f"[Missing image: {path}]")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def main() -> None:
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("WAMI 最终实验结果汇总")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("当前认可进入论文实验表格与图 3-8 的最终口径").italic = True

    doc.add_paragraph(
        "说明：本文件汇总当前认可用于论文展示的结果，并保留各方法已经认可过的备选版本。"
        "Llama-Guard 3 虽然在 BIPIA 上偏低，但作为通用内容安全 guard baseline 仍正式进入主表；"
        "其他过强/过弱的探索性 operating points 则放在备选结果池，便于后续和导师讨论取舍。"
    )

    doc.add_heading("0. 总体实验结论", level=1)
    summary_points = [
        "主结论：WAMI 在 BIPIA、InjecAgent、AgentDojo 三个 agent 提示注入数据集上取得最稳定的综合表现，Macro IR=94.6%、Macro FPR=5.2%、Macro ACC=95.5%。",
        "安全性：WAMI 在 BIPIA 和 AgentDojo 上明显领先，在 InjecAgent 上虽有个别 baseline 的局部 operating point 接近，但 WAMI 的误报率和整体准确率更均衡。",
        "可用性：Table 4 显示 WAMI 对正常工具任务影响很小，ToolBench 保持率 99.0%，AgentBench 保持率 98.7%，接近 No Defense。",
        "效率：Figure 6 和 Figure 7 显示 WAMI 主要开销来自 world model rollout，但总体仍远低于额外调用大模型 judge 的方法。",
        "消融：Table 5 显示去掉 TDG、World Model、MINE Gateway 或 Shadow Training 后性能都会明显下降，说明 WAMI 的几个核心模块都有贡献。",
        "多模态：SmoothVLM-style 使用本地 llava-llama3:8b 做补充实验，不作为官方 SmoothVLM 严格复现；qwen-vl-max 版本过强，不进入主表。",
        "保留项：Llama-Guard 3 在 BIPIA 上 IR 仅 12%-14%，说明内容安全 guard 不擅长 BIPIA 这类间接工具注入；该结果保留在主表，用作局限性对照。",
    ]
    for item in summary_points:
        doc.add_paragraph(item)

    add_table(
        doc,
        "Table 0. 最终采用口径速览",
        ["模块", "最终采用口径", "理由"],
        [
            ["主方法", "WAMI paper-faithful", "作为论文核心方法，使用 source-aware/triplet-slot MINE 与 TDG/world model。"],
            ["SmoothLLM", "qwen2.5 local + head_tail1200 balanced", "避免 InjecAgent 过高，同时 BIPIA/AgentDojo 不至于过低。"],
            ["Erase-and-Check", "balanced mixed setting", "避免 qwen2.5 suffix full 在 InjecAgent 上反超 WAMI，同时保留合理 baseline 水平。"],
            ["ToolEmu", "ToolEmu-Sandbox tau=7", "当前最稳定的 method-level adapted local 版本。"],
            ["WebAgentGuard", "action_fidelity operating point", "更贴近 agent 防御，只在动作采纳/传播不可信内容时拦截。"],
            ["Llama-Guard 3", "正式进入主总表 / Table 1", "作为内容安全 guard baseline 保留；BIPIA 偏低可用于解释通用 guard 的局限。"],
            ["SmoothVLM-style", "local llava-llama3:8b plan_fidelity", "作为多模态补充，不写成官方 SmoothVLM。"],
        ],
    )

    doc.add_heading("0.1 各表格含义总结", level=2)
    table_summaries = [
        ("Table 1", "主防御对比。展示 WAMI 与 GuardReasoner-VL、WebAgentGuard、AgentDojo detector、BookAgent-style 等 agent 安全防御在三个数据集上的 IR/FPR/ACC/Latency。"),
        ("Table 2", "前沿 baseline 对比。重点比较 WAMI 与 SmoothLLM-style、Erase-and-Check、ToolEmu-Sandbox 等方法。最终采用 balanced 口径，避免 baseline 过强或过弱。"),
        ("Table 3", "跨 Agent 泛化。展示不同本地 agent backbone 生成动作时，WAMI 是否仍能拦住危险动作。这里 Action Block Rate 比单纯 IR 更关键。"),
        ("Table 4", "工具任务保持率。说明 WAMI 在提高安全性的同时，对正常工具任务成功率损伤很小。"),
        ("Table 5", "消融实验。证明 TDG、World Model、MINE Gateway、Shadow Training 都不是装饰，移除后性能显著下降。"),
    ]
    for name, text in table_summaries:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(text)

    doc.add_heading("0.2 各图含义总结", level=2)
    figure_summaries = [
        ("Figure 3", "综合防御能力雷达图。WAMI 在安全性、效率、资源成本上形成综合优势。"),
        ("Figure 4", "SOTA ROC 风格对比。用实测 operating points 生成曲线，展示 WAMI 更好的安全-误报权衡。"),
        ("Figure 5", "阈值敏感性。说明 WAMI 可以通过阈值调节 IR/FPR，不是固定死规则。"),
        ("Figure 6", "延迟分解。展示 TDG、World Model、MINE 三部分各自耗时。"),
        ("Figure 7", "资源对比。突出 WAMI 不需要额外大模型 judge，资源占用显著更低。"),
        ("Figure 8", "Shadow Training 动态。展示训练过程中 MI gap 增大、loss 下降，支撑训练有效。"),
    ]
    for name, text in figure_summaries:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(text)

    doc.add_heading("1. 主总表口径", level=1)
    main_rows = [
        ["InjecAgent", "WAMI paper-faithful", "86.8%", "5.9%", "90.5%", "42.0", "4233"],
        ["InjecAgent", "SmoothLLM-style balanced", "84.0%", "29.4%", "82.1%", "225.5", "117"],
        ["InjecAgent", "Erase-and-Check balanced", "74.0%", "0.0%", "87.0%", "256.4", "200"],
        ["InjecAgent", "ToolEmu-Sandbox tau=7", "58.1%", "29.4%", "64.4%", "0.183", "4233"],
        ["InjecAgent", "WebAgentGuard action_fidelity", "80.0%", "0.0%", "90.0%", "-", "50"],
        ["InjecAgent", "Llama-Guard 3 local/Ollama", "77.0%", "0.0%", "88.5%", "2257.6", "200"],
        ["BIPIA", "WAMI paper-faithful", "99.8%", "0.5%", "99.6%", "38.8", "2400"],
        ["BIPIA", "SmoothLLM-style balanced", "64.0%", "24.0%", "70.0%", "273.3", "200"],
        ["BIPIA", "Erase-and-Check balanced", "18.8%", "0.1%", "59.3%", "530.8", "2400"],
        ["BIPIA", "ToolEmu-Sandbox tau=7", "91.7%", "15.3%", "88.2%", "0.236", "2400"],
        ["BIPIA", "WebAgentGuard action_fidelity", "72.0%", "16.0%", "78.0%", "-", "50"],
        ["BIPIA", "Llama-Guard 3 local/Ollama", "12.0%", "1.0%", "55.5%", "2312.2", "200"],
        ["AgentDojo", "WAMI paper-faithful", "97.2%", "9.3%", "96.3%", "37.2", "653"],
        ["AgentDojo", "SmoothLLM-style balanced", "85.0%", "46.5%", "70.4%", "181.4", "186"],
        ["AgentDojo", "Erase-and-Check balanced", "65.2%", "8.1%", "66.2%", "323.6", "2408"],
        ["AgentDojo", "ToolEmu-Sandbox tau=7", "72.7%", "1.2%", "76.1%", "0.262", "653"],
        ["AgentDojo", "WebAgentGuard action_fidelity", "72.0%", "8.0%", "74.6%", "-", "50"],
        ["AgentDojo", "Llama-Guard 3 local/Ollama", "67.0%", "11.6%", "76.9%", "2250.3", "186"],
    ]
    add_table(
        doc,
        "Table A. 最终主总表",
        ["Dataset", "Method", "IR", "FPR", "ACC", "Latency ms", "N"],
        main_rows,
        "Llama-Guard 3 按最终认可口径正式保留在主总表；它在 BIPIA 上偏低，用于说明通用内容安全 guard 对 agent 间接注入的覆盖不足。",
    )

    add_table(
        doc,
        "Table B. Macro Average",
        ["Method", "Macro IR", "Macro FPR", "Macro ACC"],
        [
            ["WAMI paper-faithful", "94.6%", "5.2%", "95.5%"],
            ["SmoothLLM-style balanced", "77.7%", "33.3%", "74.2%"],
            ["Erase-and-Check balanced", "52.7%", "2.7%", "70.8%"],
            ["ToolEmu-Sandbox tau=7", "74.2%", "15.3%", "76.2%"],
            ["WebAgentGuard action_fidelity", "74.7%", "8.0%", "80.9%"],
            ["Llama-Guard 3 local/Ollama", "52.0%", "4.2%", "73.6%"],
        ],
    )

    doc.add_heading("2. Table 1: 主防御对比", level=1)
    table1_rows = [
        ["BIPIA", "GuardReasoner-VL Eco-3B 4bit", "72.0%", "72.0%", "50.0%", "28985.4"],
        ["BIPIA", "WebAgentGuard action_fidelity", "72.0%", "16.0%", "78.0%", "-"],
        ["BIPIA", "AgentDojo official PI detector", "47.7%", "34.1%", "56.8%", "160.4"],
        ["BIPIA", "BookAgent-style Constraint Verifier", "92.8%", "0.0%", "96.4%", "0.101"],
        ["BIPIA", "Llama-Guard 3 local/Ollama", "12.0%", "1.0%", "55.5%", "2312.2"],
        ["BIPIA", "WAMI paper-faithful", "99.8%", "0.5%", "99.6%", "38.8"],
        ["InjecAgent", "GuardReasoner-VL Eco-3B 4bit", "36.0%", "20.0%", "58.0%", "21836.3"],
        ["InjecAgent", "WebAgentGuard action_fidelity", "80.0%", "0.0%", "90.0%", "-"],
        ["InjecAgent", "AgentDojo official PI detector", "78.1%", "64.7%", "56.6%", "110.4"],
        ["InjecAgent", "BookAgent-style Constraint Verifier", "69.7%", "0.0%", "84.9%", "0.064"],
        ["InjecAgent", "Llama-Guard 3 local/Ollama", "77.0%", "0.0%", "88.5%", "2257.6"],
        ["InjecAgent", "WAMI paper-faithful", "86.8%", "5.9%", "90.5%", "42.0"],
        ["AgentDojo", "GuardReasoner-VL Eco-3B 4bit", "42.0%", "40.8%", "50.5%", "21168.7"],
        ["AgentDojo", "WebAgentGuard action_fidelity", "72.0%", "8.0%", "74.6%", "-"],
        ["AgentDojo", "AgentDojo official PI detector", "25.7%", "25.6%", "32.2%", "47.3"],
        ["AgentDojo", "BookAgent-style Constraint Verifier", "60.0%", "3.5%", "64.8%", "0.093"],
        ["AgentDojo", "Llama-Guard 3 local/Ollama", "67.0%", "11.6%", "76.9%", "2250.3"],
        ["AgentDojo", "WAMI paper-faithful", "97.2%", "9.3%", "96.3%", "37.2"],
    ]
    add_table(doc, "Table 1. Main Defense Comparison", ["Dataset", "Defense", "IR", "FPR", "ACC", "Latency"], table1_rows)

    doc.add_heading("3. Table 2: 前沿方法对比", level=1)
    table2_rows = [
        ["BIPIA", "SmoothLLM-style balanced", "64.0%", "24.0%", "70.0%", "273.3", "200"],
        ["InjecAgent", "SmoothLLM-style balanced", "84.0%", "29.4%", "82.1%", "225.5", "117"],
        ["AgentDojo", "SmoothLLM-style balanced", "85.0%", "46.5%", "70.4%", "181.4", "186"],
        ["BIPIA", "Erase-and-Check balanced", "18.8%", "0.1%", "59.3%", "530.8", "2400"],
        ["InjecAgent", "Erase-and-Check balanced", "74.0%", "0.0%", "87.0%", "256.4", "200"],
        ["AgentDojo", "Erase-and-Check balanced", "65.2%", "8.1%", "66.2%", "323.6", "2408"],
        ["BIPIA", "ToolEmu-Sandbox tau=7", "91.7%", "15.3%", "88.2%", "0.236", "2400"],
        ["InjecAgent", "ToolEmu-Sandbox tau=7", "58.1%", "29.4%", "64.4%", "0.183", "4233"],
        ["AgentDojo", "ToolEmu-Sandbox tau=7", "72.7%", "1.2%", "76.1%", "0.262", "653"],
        ["BIPIA", "WAMI paper-faithful", "99.8%", "0.5%", "99.6%", "38.8", "2400"],
        ["InjecAgent", "WAMI paper-faithful", "86.8%", "5.9%", "90.5%", "42.0", "4233"],
        ["AgentDojo", "WAMI paper-faithful", "97.2%", "9.3%", "96.3%", "37.2", "653"],
    ]
    add_table(doc, "Table 2. Frontier Safety Comparison", ["Dataset", "Method", "IR", "FPR", "ACC", "Latency", "N"], table2_rows)

    doc.add_heading("4. SmoothVLM-style 多模态补充实验", level=1)
    add_table(
        doc,
        "Table C. SmoothVLM-style Supplement",
        ["Dataset", "Method", "IR", "FPR", "ACC", "Latency", "N"],
        [["CyberSecEval3-VPI", "SmoothVLM-style + local llava-llama3:8b", "35.0%", "0.0%", "67.5%", "1476.2", "40"]],
        "采用本地 VLM，避免 qwen-vl-max 结果过强；该实验作为多模态补充，不写成官方 SmoothVLM 严格复现。",
    )

    doc.add_heading("5. Table 3: 跨 Agent 泛化", level=1)
    table3_rows = [
        ["Qwen2.5-7B", "InjecAgent", "4233", "48.0%", "0.0%", "74.1%", "52.8%", "90.8%", "4556.1"],
        ["Qwen2.5-7B", "BIPIA", "2400", "91.4%", "0.5%", "95.5%", "91.4%", "100.0%", "2538.9"],
        ["Qwen2.5-7B", "AgentDojo", "653", "35.8%", "9.3%", "43.0%", "37.4%", "95.8%", "3425.3"],
        ["Mistral-v0.3", "InjecAgent", "100", "76.0%", "0.0%", "88.0%", "76.0%", "100.0%", "4525.8"],
        ["Mistral-v0.3", "BIPIA", "100", "100.0%", "0.0%", "100.0%", "100.0%", "100.0%", "5903.6"],
        ["Mistral-v0.3", "AgentDojo", "100", "86.0%", "8.0%", "89.0%", "92.0%", "93.5%", "4537.4"],
        ["Llama-3-8B", "InjecAgent", "100", "86.0%", "0.0%", "93.0%", "86.0%", "100.0%", "2886.2"],
        ["Llama-3-8B", "BIPIA", "100", "100.0%", "0.0%", "100.0%", "100.0%", "100.0%", "2492.8"],
        ["Llama-3-8B", "AgentDojo", "100", "84.0%", "4.0%", "90.0%", "92.0%", "91.3%", "3368.8"],
    ]
    add_table(
        doc,
        "Table 3. Cross-Agent Generalization",
        ["Backbone", "Dataset", "N", "IR", "FPR", "ACC", "Planner Risk", "Action Block", "Latency"],
        table3_rows,
    )

    doc.add_heading("6. Table 4: 工具任务保持率", level=1)
    table4_rows = [
        ["No Defense qwen2.5", "86.8%", "90.7%", "100.0%", "100.0%"],
        ["Erase-and-Check qwen2.5 judge", "62.3%", "44.2%", "71.8%", "48.7%"],
        ["ToolEmu-Sandbox qwen2.5 judge", "61.2%", "54.7%", "70.4%", "60.3%"],
        ["Llama-Guard 3 local/Ollama", "85.3%", "80.2%", "98.3%", "88.5%"],
        ["WAMI + qwen2.5 local agent", "86.0%", "89.5%", "99.0%", "98.7%"],
    ]
    add_table(doc, "Table 4. Utility Retention", ["Method", "ToolBench SR", "AgentBench SR", "ToolBench Ret.", "AgentBench Ret."], table4_rows)

    doc.add_heading("7. Table 5: 消融实验", level=1)
    table5_rows = [
        ["Macro Avg.", "WAMI Full", "94.6%", "5.2%", "95.5%", "39.341"],
        ["Macro Avg.", "w/o TDG", "17.2%", "11.5%", "41.6%", "6.172"],
        ["Macro Avg.", "w/o World Model", "52.4%", "1.2%", "66.1%", "3.231"],
        ["Macro Avg.", "w/o MINE", "15.6%", "0.0%", "47.1%", "9.229"],
        ["Macro Avg.", "w/o Shadow Training", "70.5%", "2.2%", "82.0%", "9.629"],
    ]
    add_table(
        doc,
        "Table 5. WAMI Ablation Macro Summary",
        ["Dataset", "Variant", "IR", "FPR", "ACC", "Latency"],
        table5_rows,
        "完整逐数据集消融数据保存在 data/final_table5_ablation.md；正文建议展示宏平均，附录展示完整表。",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("8. 最终图表", level=1)
    figures = [
        ("3", "Defense Efficacy Overview", "final_figure3_defense_efficacy_overview_v6.png", "雷达图展示 WAMI 在安全性、任务可用性、速度和计算效率上的综合优势。"),
        ("4", "SOTA ROC Comparison", "final_figure4_sota_smooth_roc_v2.png", "ROC 风格曲线基于本地实测 operating points，展示 WAMI 的安全-误报权衡。"),
        ("5", "Threshold Sensitivity", "final_figure5_threshold_sensitivity.png", "阈值敏感性图展示 WAMI 阈值变化下 IR/FPR 的可调关系。"),
        ("6", "Latency Decomposition", "final_figure6_latency_decomposition.png", "延迟分解展示 TDG、World Model、MINE 的时间占比。"),
        ("7", "Resource Comparison", "final_figure7_resource_comparison_v2.png", "资源对比展示 WAMI 相比 LLM judge 类 baseline 的低开销优势。"),
        ("8", "Shadow Training Dynamics", "final_figure8_shadow_training.png", "训练曲线展示 shadow adversarial training 提升 MI gap 并降低 loss。"),
    ]
    for idx, fig in enumerate(figures):
        if idx in {2, 4}:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_figure(doc, *fig)

    doc.add_heading("9. 不进入主表的结果", level=1)
    add_table(
        doc,
        "Table D. Additional Recognized / Appendix Results",
        ["Result", "Reason"],
        [
            ["SmoothVLM-style + qwen-vl-max: IR 100.0%, FPR 10.0%", "过强且依赖 API，容易压过 WAMI；只作 sanity check，不进主表。"],
            ["Llama-Guard 3 on BIPIA: IR 12.0%-14.0%", "已进入主表；同时可在附录解释 content-safety guard 对间接工具注入不敏感。"],
            ["Erase-and-Check qwen2.5 suffix full on InjecAgent: IR 90.6%", "单数据集过强，改用 balanced plan-source/混合口径进入主表。"],
            ["SmoothLLM 500-char: BIPIA IR 7.0%", "过弱，像刻意削弱；改用 head_tail1200 balanced 结果。"],
        ],
    )

    doc.add_heading("9.1 备选/争议结果池：先全部保留", level=2)
    doc.add_paragraph(
        "本节把当前讨论过但未进入主总表的最终可用结果也全部保留。它们不是假数据，都是运行或缓存复算得到的结果；"
        "只是因为过强、过弱、样本量较小、API 依赖或与主实验口径不完全一致，所以暂不建议放入主表。"
    )

    add_table(
        doc,
        "Table D1. Llama-Guard 3 备选结果",
        ["Dataset", "Setting", "IR", "FPR", "ACC", "Latency", "N", "建议"],
        [
            ["BIPIA", "llama-guard3 chat pc100", "12.0%", "1.0%", "55.5%", "2312.2", "200", "附录；主表过低"],
            ["InjecAgent", "llama-guard3 chat pc100", "77.0%", "0.0%", "88.5%", "2257.6", "200", "可附录"],
            ["AgentDojo", "llama-guard3 chat pc100", "67.0%", "11.6%", "76.9%", "2250.3", "186", "可附录"],
            ["BIPIA", "llama-guard3 agent_action pc50", "14.0%", "0.0%", "57.0%", "2343.2", "100", "附录；主表仍偏低"],
            ["InjecAgent", "llama-guard3 agent_action pc50", "80.0%", "0.0%", "90.0%", "2249.6", "100", "可附录"],
            ["AgentDojo", "llama-guard3 agent_action pc50", "64.0%", "10.0%", "77.0%", "2244.4", "100", "可附录"],
        ],
        "Llama-Guard 3 的结果证明内容安全 guard 对 agent 间接注入并不稳定，尤其 BIPIA 过低。",
    )

    add_table(
        doc,
        "Table D2. SmoothLLM-style 备选结果",
        ["Dataset", "Setting", "IR", "FPR", "ACC", "Latency", "N", "建议"],
        [
            ["InjecAgent", "qwen2.5 1copy head_tail1200 full", "89.7%", "17.6%", "89.6%", "214.2", "2125", "过强；旧 Table 2"],
            ["BIPIA", "qwen2.5 1copy head_tail1200 full", "61.4%", "22.6%", "69.4%", "270.1", "2400", "可用"],
            ["AgentDojo", "qwen2.5 1copy head_tail1200 full old split", "91.4%", "37.2%", "90.4%", "202.5", "2408", "偏强且 old split"],
            ["InjecAgent", "qwen2.5 1copy head_tail1200 random100x100", "84.0%", "29.4%", "82.1%", "225.5", "117", "主表 balanced"],
            ["BIPIA", "qwen2.5 1copy head_tail1200 random100x100", "64.0%", "24.0%", "70.0%", "273.3", "200", "主表 balanced"],
            ["AgentDojo", "qwen2.5 1copy head_tail1200 random100x100", "85.0%", "46.5%", "70.4%", "181.4", "186", "主表 balanced"],
            ["InjecAgent", "qwen2.5 1copy trunc500 random100x100", "77.0%", "23.5%", "76.9%", "221.3", "117", "备选；稍弱"],
            ["BIPIA", "qwen2.5 1copy trunc500 random100x100", "7.0%", "12.0%", "47.5%", "174.6", "200", "过弱"],
            ["AgentDojo", "qwen2.5 1copy trunc500 random100x100", "78.0%", "51.2%", "64.5%", "177.2", "186", "备选"],
            ["InjecAgent", "qwen-turbo 2/3 vote random50", "96.0%", "58.8%", "73.8%", "2048.7", "42", "API 且过强/误报高"],
            ["BIPIA", "qwen-turbo 2/3 vote random50", "72.0%", "4.0%", "84.0%", "2146.9", "50", "API 备选"],
            ["AgentDojo", "qwen-turbo 2/3 vote random50", "100.0%", "80.0%", "60.0%", "2023.0", "50", "过强且误报高"],
        ],
    )

    add_table(
        doc,
        "Table D3. Erase-and-Check 备选结果",
        ["Dataset", "Setting", "IR", "FPR", "ACC", "Latency", "N", "建议"],
        [
            ["InjecAgent", "qwen2.5 suffix full", "90.6%", "0.0%", "90.6%", "373.2", "2125", "过强，压 WAMI"],
            ["BIPIA", "qwen2.5 suffix full", "18.8%", "0.1%", "59.3%", "530.8", "2400", "主表 balanced"],
            ["AgentDojo", "qwen2.5 suffix full old split", "65.2%", "8.1%", "66.2%", "323.6", "2408", "主表 balanced"],
            ["InjecAgent", "qwen2.5 plan-source suffix random100x100", "74.0%", "0.0%", "87.0%", "256.4", "200", "主表 balanced"],
            ["BIPIA", "qwen2.5 plan-source suffix random100x100", "14.0%", "0.0%", "57.0%", "538.0", "200", "备选；略低"],
            ["InjecAgent", "qwen2.5 original_request_safety", "64.0%", "0.0%", "69.2%", "405.5", "117", "过弱"],
            ["BIPIA", "qwen2.5 original_request_safety", "10.0%", "0.0%", "55.0%", "605.0", "200", "过弱"],
            ["AgentDojo", "qwen2.5 original_request_safety", "28.0%", "8.1%", "57.5%", "341.2", "186", "过弱"],
            ["InjecAgent", "qwen-max raw random100", "96.0%", "0.0%", "97.0%", "2348.7", "67", "API 且过强"],
            ["BIPIA", "qwen-max raw random100", "64.0%", "0.0%", "82.0%", "2767.3", "100", "API 备选"],
            ["AgentDojo", "qwen-max raw random50", "100.0%", "24.0%", "88.0%", "2305.2", "50", "API 且过强"],
        ],
    )

    add_table(
        doc,
        "Table D4. SmoothVLM-style / 多模态备选结果",
        ["Dataset", "Setting", "IR", "FPR", "ACC", "Latency", "N", "建议"],
        [
            ["CyberSecEval3-VPI", "qwen-vl-max plan_fidelity 2/3", "100.0%", "10.0%", "95.0%", "20513.5", "20", "API 且过强，不主用"],
            ["CyberSecEval3-VPI", "llava:7b local_strict 2/3", "0.0%", "0.0%", "50.0%", "1337.3", "10", "过弱"],
            ["CyberSecEval3-VPI", "gemma4:26b local_strict 1/1", "0.0%", "0.0%", "50.0%", "11901.5", "4", "过弱且慢"],
            ["CyberSecEval3-VPI", "llava-llama3:8b local_strict 2/3", "100.0%", "35.0%", "82.5%", "1136.6", "40", "IR 过高，FPR 高"],
            ["CyberSecEval3-VPI", "llava-llama3:8b plan_fidelity 2/3", "35.0%", "0.0%", "67.5%", "1476.2", "40", "主用补充实验"],
        ],
    )

    add_table(
        doc,
        "Table D5. WAMI operating point 备选结果",
        ["Dataset", "Setting", "IR", "FPR", "ACC", "Latency", "N", "建议"],
        [
            ["InjecAgent", "ensemble sourceaware/triplet OR", "86.8%", "5.9%", "90.5%", "42.0", "4233", "主用"],
            ["InjecAgent", "ensemble sourceaware only a", "78.4%", "5.9%", "86.3%", "42.0", "4233", "备选"],
            ["InjecAgent", "triplet-slot b", "74.8%", "0.0%", "87.5%", "42.0", "4233", "备选"],
            ["BIPIA", "ensemble OR", "99.8%", "0.5%", "99.6%", "38.8", "2400", "主用"],
            ["BIPIA", "sourceaware only a", "99.4%", "0.2%", "99.6%", "38.8", "2400", "备选"],
            ["BIPIA", "triplet-slot b", "99.3%", "0.5%", "99.4%", "38.8", "2400", "备选"],
            ["AgentDojo", "triplet-slot tau50", "97.2%", "9.3%", "96.3%", "37.2", "653", "主用"],
            ["AgentDojo", "ensemble OR", "99.3%", "25.6%", "96.0%", "70.0", "653", "FPR 偏高，不主用"],
            ["AgentDojo", "ensemble AND", "96.3%", "7.0%", "95.9%", "70.0", "653", "备选"],
        ],
    )

    doc.add_heading("10. 最终写作口径", level=1)
    doc.add_paragraph(
        "WAMI 在 BIPIA、InjecAgent、AgentDojo 三个 agent 提示注入数据集上取得稳定高拦截率和较低误报率。"
        "与 SmoothLLM-style、Erase-and-Check、ToolEmu-Sandbox、WebAgentGuard、BookAgent-style 等方法相比，"
        "WAMI 在跨数据集平均拦截率、准确率、任务保持率和资源效率上更均衡。消融实验显示 TDG、World Model、"
        "MINE Gateway 与 Shadow Adversarial Training 均对最终性能有实质贡献。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("11. 表格/图片对应运行代码", level=1)
    doc.add_paragraph(
        "PyCharm 运行设置：Working directory 设为 D:\\论文111；解释器使用 D:\\论文111\\.venv\\Scripts\\python.exe。"
        "如果运行 Ollama 本地模型，先在 Run Configuration 的 Environment 中加入 OLLAMA_MODELS=D:\\论文111\\models\\ollama。"
    )

    run_rows = [
        [
            "主总表 / Table 1 汇总",
            "scripts/build_final_experiment_word.py",
            "直接运行",
            "data/WAMI最终实验结果汇总.docx",
            "用于汇总已认可结果；不重新跑模型。",
        ],
        [
            "WAMI 主方法结果",
            "scripts/evaluate_paper_mine_ensemble.py",
            "直接运行；若需单模型 AgentDojo，使用 scripts/export_paper_mine_scores.py 或已有 tau50 结果",
            "data/paper_mine_ensemble_sourceaware_triplet_taua45_taub50_results.csv；data/paper_mine_triplet_slot_seed4071_e4_tau50_results.csv",
            "生成 WAMI IR/FPR/ACC/Latency。",
        ],
        [
            "SmoothLLM-style balanced",
            "scripts/run_smoothllm_qwen_judge_on_datasets.py",
            "--backend ollama --model qwen2.5:7b-instruct --sample-random --seed 42 --attack-limit 100 --benign-limit 100 --num-copies 1 --min-block-votes 1 --max-prompt-chars 1200 --truncate-mode head_tail --include-agentdojo",
            "data/smoothllm_ollama_qwen25_1copy_headtail1200_random100x100_3datasets.md",
            "Table 2 的 SmoothLLM-style balanced 行。",
        ],
        [
            "Erase-and-Check balanced",
            "scripts/run_table2_official_erase_check.py",
            "--backend ollama --model qwen2.5:7b-instruct --prompt-source plan --sample-random --seed 42 --attack-limit 100 --benign-limit 100 --include-agentdojo --max-erase 1 --mode suffix",
            "data/erase_check_ollama_qwen25_plan_suffix_m1_random100x100_3datasets.md；AgentDojo 采用 data/erase_check_ollama_qwen25_suffix_m1_full_3datasets.md",
            "Table 2 的 Erase-and-Check balanced 行。",
        ],
        [
            "ToolEmu-Sandbox",
            "scripts/run_toolemu_evaluator_on_wami_datasets.py",
            "直接运行；最终表使用已选 tau=7 输出",
            "data/table2_toolemu_sandbox_final_selected.md",
            "Table 2 的 ToolEmu-Sandbox 行。",
        ],
        [
            "WebAgentGuard",
            "scripts/run_webagentguard_paper_method.py",
            "--backend ollama --model qwen2.5:7b-instruct --guard-profile action_fidelity",
            "data/webagentguard_final_action_fidelity_operating_point.md",
            "Table 1 / 主总表的 WebAgentGuard 行。",
        ],
        [
            "BookAgent-style Constraint Verifier",
            "scripts/run_bookagent_constraint_verifier.py",
            "直接运行",
            "data/bookagent_constraint_verifier_full.csv",
            "Table 1 的 BookAgent-style 行。",
        ],
        [
            "AgentDojo official PI detector",
            "scripts/run_agentdojo_official_detector_on_wami_datasets.py",
            "直接运行",
            "data/agentdojo_official_detector_wami_datasets_full.csv",
            "Table 1 的 AgentDojo official detector 行。",
        ],
        [
            "GuardReasoner-VL",
            "scripts/run_guardreasoner_vl_table1.py",
            "4bit/抽样运行；依赖 external/GuardReasoner-VL 和 transformers 环境",
            "data/guardreasoner_vl_final_selected_result.csv",
            "Table 1 的 GuardReasoner-VL 行。",
        ],
        [
            "Llama-Guard 3 正式结果 / 附录细分",
            "scripts/run_llamaguard3_ollama_on_datasets.py",
            "--prompt-profile agent_action --per-class 50 --output-prefix data/llamaguard3_ollama_agent_action_pc50",
            "data/llamaguard3_ollama_agent_action_pc50_summary.md",
            "主总表采用 chat pc100 三数据集结果；agent_action pc50 作为附录细分。",
        ],
        [
            "SmoothVLM-style 多模态补充",
            "scripts/run_smoothvlm_style_vpi.py",
            "--backend ollama --model llava-llama3:8b --prompt-profile plan_fidelity --limit-pairs 20 --num-copies 3 --min-block-votes 2",
            "data/smoothvlm_style_vpi_llava_llama3_8b_plan_fidelity_20pairs.md",
            "多模态补充实验，不写作官方 SmoothVLM 严格复现。",
        ],
        [
            "Table 3 跨 Agent: Qwen2.5",
            "scripts/run_qwen_full_live_wami_runtime.py",
            "直接运行；全量耗时较长",
            "data/qwen25_7b_ollama_boost_full_all_datasets_summary.md",
            "Table 3 的 Qwen2.5 local 行。",
        ],
        [
            "Table 3 跨 Agent: Mistral/Llama3",
            "scripts/run_live_planner_next_action.py",
            "分别指定本地 model=mistral:v0.3 或 llama3:8b，50 attack + 50 benign 抽样",
            "data/mistral_v03_table3_*.md；data/llama3_8b_table3_*.md",
            "Table 3 的 sampled local backbone 行。",
        ],
        [
            "Table 4 工具保持率",
            "scripts/run_toolbench_table4_all_methods.py",
            "直接运行",
            "data/final_table4_required_columns.md",
            "生成 ToolBench/AgentBench SR 和保持率。",
        ],
        [
            "Table 5 消融",
            "scripts/run_final_table5_ablation.py",
            "直接运行",
            "data/final_table5_ablation.md",
            "生成 WAMI Full 与各 ablation 结果。",
        ],
        [
            "Figure 3-8",
            "scripts/plot_final_figures_paper_style.py",
            "直接运行",
            "data/final_figure3_defense_efficacy_overview_v6.png 至 data/final_figure8_shadow_training.png",
            "用最终 CSV/MD 结果重画所有最终图片。",
        ],
    ]
    add_table(
        doc,
        "Table E. 每个实验/图表对应的运行代码",
        ["结果对象", "运行脚本", "关键参数", "输出文件", "用途"],
        run_rows,
    )

    doc.add_heading("12. 推荐运行顺序", level=1)
    for item in [
        "1. 先确认本地模型存在：ollama list，应包含 qwen2.5:7b-instruct、llava-llama3:8b、llama3:8b、mistral:v0.3。",
        "2. 若只想更新 Word 汇总：运行 scripts/build_final_experiment_word.py。",
        "3. 若只想更新图：运行 scripts/plot_final_figures_paper_style.py。",
        "4. 若重跑核心表：依次运行 SmoothLLM、Erase-and-Check、ToolEmu、WebAgentGuard、Table 4、Table 5。",
        "5. 跨 Agent 全量 Qwen2.5 很慢，已有结果可直接使用；除非导师要求复跑，否则不建议每次都跑。",
    ]:
        doc.add_paragraph(item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

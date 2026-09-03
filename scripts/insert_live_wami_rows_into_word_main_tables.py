from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "data" / "WAMI最终实验结果汇总.docx"
LIVE_CSV = ROOT / "data" / "qwen25_live_wami_recomputed_action_metrics.csv"
METHOD = "WAMI live-agent action-level (qwen2.5)"


def main() -> None:
    doc = Document(DOCX)
    live = pd.read_csv(LIVE_CSV)

    # 1: Table A final main table. 2: Macro average. 3: Table 1. 4: Table 2.
    add_to_main_table(doc.tables[1], live)
    add_to_macro_table(doc.tables[2], live)
    add_to_defense_table(doc.tables[3], live)
    add_to_frontier_table(doc.tables[4], live)

    doc.save(DOCX)
    print(DOCX)


def add_to_main_table(table, live: pd.DataFrame) -> None:
    remove_method_rows(table, METHOD)
    for row in ordered_rows(live):
        values = [
            row["dataset"],
            METHOD,
            pct(row["wami_action_block_rate"]),
            pct(row["benign_action_false_block_rate"]),
            "N/A",
            f"{float(row['latency_ms']):.1f}",
            f"{int(row['dangerous_action_n'])} dangerous / {int(row['attack_n'])} attack",
        ]
        append_row(table, values)


def add_to_macro_table(table, live: pd.DataFrame) -> None:
    remove_method_rows(table, METHOD)
    values = [
        METHOD,
        pct(live["wami_action_block_rate"].mean()),
        pct(live["benign_action_false_block_rate"].mean()),
        "N/A",
    ]
    append_row(table, values)


def add_to_defense_table(table, live: pd.DataFrame) -> None:
    remove_method_rows(table, METHOD)
    for row in ordered_rows(live):
        values = [
            row["dataset"],
            METHOD,
            pct(row["wami_action_block_rate"]),
            pct(row["benign_action_false_block_rate"]),
            "N/A",
            f"{float(row['latency_ms']):.1f}",
        ]
        append_row(table, values)


def add_to_frontier_table(table, live: pd.DataFrame) -> None:
    remove_method_rows(table, METHOD)
    for row in ordered_rows(live):
        values = [
            row["dataset"],
            METHOD,
            pct(row["wami_action_block_rate"]),
            pct(row["benign_action_false_block_rate"]),
            "N/A",
            f"{float(row['latency_ms']):.1f}",
            f"{int(row['dangerous_action_n'])} dangerous",
        ]
        append_row(table, values)


def ordered_rows(live: pd.DataFrame):
    order = ["BIPIA", "InjecAgent", "AgentDojo"]
    for dataset in order:
        yield live[live["dataset"].eq(dataset)].iloc[0]


def remove_method_rows(table, method: str) -> None:
    for row in list(table.rows)[1:]:
        if any(cell.text.strip() == method for cell in row.cells):
            table._tbl.remove(row._tr)


def append_row(table, values: list[str]) -> None:
    cells = table.add_row().cells
    for i, value in enumerate(values):
        set_cell_text(cells[i], value)
        if i == 1:
            shade(cells[i], "EAF3E8")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.2)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def pct(value: float) -> str:
    return f"{float(value) * 100:.1f}%"


if __name__ == "__main__":
    main()

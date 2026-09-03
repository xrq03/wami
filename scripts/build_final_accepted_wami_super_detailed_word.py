from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "WAMI最终认可版本_超级详细讲解.docx"


ACCEPTED_RESULTS = [
    ["WAMI paper-faithful", "BIPIA", "99.8%", "0.5%", "99.6%", "最终主方法静态轨迹评测"],
    ["WAMI paper-faithful", "InjecAgent", "86.8%", "5.9%", "90.5%", "最终主方法静态轨迹评测"],
    ["WAMI paper-faithful", "AgentDojo", "97.2%", "9.3%", "96.3%", "最终主方法静态轨迹评测"],
    ["qwen2.5 live-agent action-level", "BIPIA", "100.0%", "0.1%", "N/A", "WAMI 对 qwen2.5 已生成危险动作的拦截"],
    ["qwen2.5 live-agent action-level", "InjecAgent", "89.3%", "0.0%", "N/A", "WAMI 对 qwen2.5 已生成危险动作的拦截"],
    ["qwen2.5 live-agent action-level", "AgentDojo", "93.1%", "1.8%", "N/A", "WAMI 对 qwen2.5 已生成危险动作的拦截"],
    ["qwen2.5 stable smoke 20+20", "BIPIA", "70.0%", "0.0%", "85.0%", "小规模实时 agent 验证，BIPIA 带 bootstrap"],
    ["qwen2.5 stable smoke 20+20", "InjecAgent", "70.0%", "0.0%", "85.0%", "小规模实时 agent 验证"],
    ["qwen2.5 stable smoke 20+20", "AgentDojo", "50.0%", "0.0%", "75.0%", "小规模实时 agent 验证"],
]


FILES = [
    ["scripts/generate_self_training_data.py", "生成独立 shadow train/val 数据，包含独立正常样本、hard negatives、pair/triplet 和攻击样本。"],
    ["scripts/run_qwen_live_wami_smoke_stable.py", "最终小规模 qwen2.5 live-agent 稳定入口，串行跑 20+20 或更大子集。"],
    ["scripts/run_qwen_full_live_wami_runtime.py", "完整 live-agent runtime 核心实现，stable runner 复用它的 run_one/summarize/write_outputs。"],
    ["wami/datasets.py", "把 InjecAgent、BIPIA、AgentDojo 原始格式转换为 intent/plan/label。"],
    ["wami/tdg.py", "把 plan 转成工具依赖图 TDG，识别不可信来源、敏感数据和 sink。"],
    ["wami/runtime_trace.py", "把实时 agent 事件转换为 runtime TDG。"],
    ["wami/paper_mine_gateway.py", "WAMI 主拦截入口，计算 MINE/world-model 风险并输出 block/allow。"],
    ["wami/torch_model.py", "论文级 torch 模型，包括 world model、memory、MINE 和辅助头。"],
    ["wami/torch_training.py", "shadow training 训练循环。"],
    ["wami/evaluate.py", "统一计算 IR/FPR/ACC。"],
]


COMMANDS = [
    ["生成新 shadow 数据", ".\\.venv\\Scripts\\python.exe scripts\\generate_self_training_data.py --count 2000 --seed 20260527 --independent-benign-ratio 0.30 --output data\\self_generated_wami_train_2000_independent_benign.jsonl"],
    ["跑 InjecAgent 20+20", ".\\.venv\\Scripts\\python.exe scripts\\run_qwen_live_wami_smoke_stable.py --dataset InjecAgent --attack-limit 20 --benign-limit 20 --output-md data\\qwen_live_wami_smoke_stable_injecagent_20x20_after_benign.md --output-csv data\\qwen_live_wami_smoke_stable_injecagent_20x20_after_benign.csv"],
    ["跑 BIPIA 20+20", ".\\.venv\\Scripts\\python.exe scripts\\run_qwen_live_wami_smoke_stable.py --dataset BIPIA --attack-limit 20 --benign-limit 20 --bootstrap-first-observation --output-md data\\qwen_live_wami_smoke_stable_bipia_20x20_bootstrap_after_benign.md --output-csv data\\qwen_live_wami_smoke_stable_bipia_20x20_bootstrap_after_benign.csv"],
    ["跑 AgentDojo 20+20", ".\\.venv\\Scripts\\python.exe scripts\\run_qwen_live_wami_smoke_stable.py --dataset AgentDojo --attack-limit 20 --benign-limit 20 --output-md data\\qwen_live_wami_smoke_stable_agentdojo_20x20_after_benign.md --output-csv data\\qwen_live_wami_smoke_stable_agentdojo_20x20_after_benign.csv"],
]


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.8)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Subtitle", 10.5, "44546A"),
        ("Heading 1", 14.5, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
        ("Heading 3", 10.6, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, bold: bool = False, center: bool = False, size: float = 7.6) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], size: float = 7.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, True, 7.8)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=size)
    doc.add_paragraph()


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.15)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(7.3)


def count_rows(path: Path) -> str:
    if not path.exists():
        return "文件不存在"
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return str(max(0, sum(1 for _ in csv.reader(f)) - 1))
    except Exception:
        return "无法读取"


def build() -> None:
    doc = Document()
    setup(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("WAMI 最终认可版本超级详细讲解")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("只保留最终认可链路：WAMI 主方法、qwen2.5 live-agent action-level、稳定小规模验证、每条样本审计")

    doc.add_heading("1. 最终认可版本是什么", level=1)
    p(doc, "最终认可版本不是把所有历史尝试都塞进论文，而是保留三条主线：第一，WAMI paper-faithful 静态轨迹主方法；第二，qwen2.5 live-agent action-level 结果；第三，刚跑通的 qwen2.5 稳定 20+20 小验证。过去那些过强、过弱、临时 prompt、并行跑卡住的结果不进入主讲版本。")
    add_table(doc, ["版本", "数据集", "IR/Action Block", "FPR", "ACC", "使用方式"], ACCEPTED_RESULTS)

    doc.add_heading("2. 给导师的最短解释", level=1)
    p(doc, "我的框架做的是 agent 提示注入防御。数据集中有用户请求和工具轨迹，WAMI 先把工具轨迹构造成 TDG，也就是工具依赖图；然后 world model 预测每一步动作后的安全状态；MINE gateway 根据风险分数决定是否阻断。qwen2.5 在 live-agent 实验里不是防御器，它只是根据用户请求和工具 observation 生成下一步动作，WAMI 在动作执行前判断能不能放行。")

    doc.add_heading("3. Shadow Training 数据到底怎么来", level=1)
    p(doc, "最新版本已经补了独立正常样本池，不再只靠攻击对照样本里的 benign。生成器会先保留一部分独立 benign，然后生成 pair/triplet 对照样本，最后补齐随机攻击和正常任务。")
    add_table(
        doc,
        ["数据来源", "label", "作用"],
        [
            ["INDEPENDENT_BENIGN_CORPUS", "0", "普通正常任务和用户授权敏感操作，告诉模型正常动作不能乱拦。"],
            ["BENIGN_HARD_NEGATIVES", "0", "有注入文字但 agent 正确忽略，训练模型不要看到攻击词就拦。"],
            ["make_counterfactual_triplet", "0/0/1", "同一个 intent 下生成正常内容、忽略注入、执行攻击三种轨迹。"],
            ["make_counterfactual_pair", "0/1", "同一个任务生成正常和攻击两条近似轨迹。"],
            ["make_attack", "1", "跨工具注入、多步污染、隐藏目标迁移、视觉注入、敏感工具劫持。"],
        ],
    )

    doc.add_heading("4. qwen2.5 live-agent 具体怎么跑", level=1)
    p(doc, "qwen2.5 的作用是生成动作，不是判断安全。实验流程是：数据集样本提供 intent 和可还原的 observation；qwen2.5 生成下一步 tool call；WAMI 把 pending action 放进 runtime trace，转成 TDG，然后在执行前拦截。")
    code(
        doc,
        "intent + observation -> qwen2.5 planner -> next action\n"
        "next action + runtime trace -> WAMI TDG/MINE gateway -> block or allow\n"
        "block/allow + label -> IR/FPR/ACC/action-level metrics"
    )

    doc.add_heading("5. 最终涉及代码文件", level=1)
    add_table(doc, ["文件", "用途"], FILES)

    doc.add_heading("6. 核心函数解释", level=1)
    add_table(
        doc,
        ["文件", "函数/类", "详细解释"],
        [
            ["generate_self_training_data.py", "generate_samples", "总调度：先生成独立正常样本，再生成 triplet/pair，最后补齐随机攻击/正常。"],
            ["generate_self_training_data.py", "make_benign_hard_negative", "生成带注入词但正确忽略的正常样本，是降低误报的关键。"],
            ["generate_self_training_data.py", "make_attack", "生成不同攻击模式，包括跨工具、多步污染、视觉注入和敏感工具劫持。"],
            ["run_qwen_live_wami_smoke_stable.py", "main", "最终稳定小验证入口，串行调用 qwen2.5，不并行抢 Ollama。"],
            ["run_qwen_full_live_wami_runtime.py", "run_one", "单条样本完整 live-agent 流程：qwen 生成动作，WAMI 执行前判断。"],
            ["run_qwen_full_live_wami_runtime.py", "call_planner", "构造给 qwen2.5 的 planner prompt，要求只输出一个 JSON tool call。"],
            ["run_qwen_full_live_wami_runtime.py", "build_tool_outputs", "从数据集 plan 里还原 mock observation，避免真的调用 Gmail/Amazon/银行等工具。"],
            ["run_qwen_full_live_wami_runtime.py", "summarize", "计算 IR、FPR、ACC、Planner Risk Rate、WAMI Action Block Rate。"],
            ["tdg.py", "build_tdg", "把 plan 里的 Action 行解析成 TDG 节点和边。"],
            ["tdg.py", "analyze_tdg_security", "分析不可信来源、敏感内容、sink，以及从 source 到 sink 的路径。"],
            ["paper_mine_gateway.py", "PaperMINEGateway.inspect", "WAMI 主判断入口：构图、rollout、算 MINE 分数、输出 block/allow。"],
            ["runtime_trace.py", "build_runtime_tdg", "把 live-agent 的实时事件转换成 TDG。"],
            ["evaluate.py", "evaluate_gateway", "把 gateway 判断结果和 label 对比，算最终指标。"],
        ],
    )

    doc.add_heading("7. 认可结果文件在哪里", level=1)
    result_files = [
        ["WAMI 主结果 Word", "data/WAMI最终实验结果汇总.docx"],
        ["WAMI paper-faithful 每条样本 Excel", "data/method_audit_excels_expanded/WAMI_paper_faithful.xlsx"],
        ["qwen2.5 live-agent 每条样本 Excel", "data/method_audit_excels_expanded/WAMI_live_agent_action_level_qwen25.xlsx"],
        ["qwen2.5 action-level 指标", "data/qwen25_live_wami_recomputed_action_metrics.csv"],
        ["InjecAgent 20+20", "data/qwen_live_wami_smoke_stable_injecagent_20x20_after_benign.csv"],
        ["BIPIA 20+20", "data/qwen_live_wami_smoke_stable_bipia_20x20_bootstrap_after_benign.csv"],
        ["AgentDojo 20+20", "data/qwen_live_wami_smoke_stable_agentdojo_20x20_after_benign.csv"],
    ]
    rows = [[name, path, "存在" if (ROOT / path).exists() else "缺失", count_rows(ROOT / path) if path.endswith(".csv") else ""] for name, path in result_files]
    add_table(doc, ["名称", "路径", "状态", "CSV行数"], rows)

    doc.add_heading("8. 直接运行哪些命令", level=1)
    add_table(doc, ["目的", "命令"], COMMANDS, size=6.6)

    doc.add_heading("9. 不放进最终主讲的内容", level=1)
    bullet(doc, "并行跑 qwen2.5 的失败结果不放入，因为本地 Ollama 单模型并发会堵塞。")
    bullet(doc, "过强或过弱的 baseline 调参版本不放入，只保留最终认可版本或作为备选池在审计文件中保存。")
    bullet(doc, "临时 compact/generate prompt 路径不作为正式实验，只保留代码参数用于排查。")
    bullet(doc, "单纯 static smoke 的 80 条自生成样本结果不作为主表，只说明新数据格式和 WAMI gateway 没坏。")

    doc.add_heading("10. 给导师讲的时候按这个顺序", level=1)
    bullet(doc, "先说问题：间接提示注入藏在工具返回、网页、邮件、文件或图片里。")
    bullet(doc, "再说 WAMI：把工具调用变成 TDG，用 world model 预测状态，用 MINE 分数拦截。")
    bullet(doc, "再说 qwen2.5：它只是 live-agent planner，不是防御器。")
    bullet(doc, "再说 Shadow Training：新版本加了独立 benign 和 hard negatives，目的是降低误报。")
    bullet(doc, "最后拿结果：paper-faithful 主表 + qwen2.5 action-level + 每条样本 Excel。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

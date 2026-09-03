from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "WAMI代码实验全量讲解_导师版.docx"


CORE_FILES = [
    ("wami/gateway.py", "早期/通用 WAMI gateway，封装基础拦截逻辑、分数融合和决策输出。"),
    ("wami/torch_model.py", "论文级神经网络模型：latent encoder、Transformer world model、GRU/slot memory、MINE/辅助头。"),
    ("wami/torch_training.py", "Torch 训练循环：训练 world model、MINE、source/sink/sensitive/provenance 等辅助任务。"),
    ("wami/tdg.py", "把 Action 文本解析成 TDG，分析工具依赖、不可信来源、敏感内容、sink 流向。"),
    ("wami/multimodal.py", "多模态 WAMI：图片 latent encoder、Qwen-VL/本地视觉编码、图文融合 gateway。"),
    ("wami/model.py", "轻量 NumPy/基础模型版本，用于快速实验和早期复现。"),
    ("wami/datasets.py", "官方数据集转换器：InjecAgent/BIPIA/AgentDojo -> intent/plan/label。"),
    ("wami/runtime_trace.py", "工业级 live runtime trace：记录用户输入、工具调用、observation、assistant action 并转 TDG。"),
    ("wami/live_agent.py", "Qwen/Ollama 等 live agent planner：让模型根据 intent+observation 生成下一步 action。"),
    ("wami/shadow_llm.py", "LLM 生成 shadow attack：跨工具注入、多步间接注入、上下文污染、视觉注入等。"),
    ("wami/paper_mine_gateway.py", "最终 paper-faithful MINE gateway：调用 TDG+world model+MINE score 做 block/allow。"),
    ("wami/shadow.py", "非 LLM shadow attack 和合成样本生成。"),
    ("wami/baselines.py", "轻量 baseline：Keyword、Erase-and-Check Lite、SmoothVLM Lite、ToolEmu Lite 等。"),
    ("wami/training.py", "轻量训练逻辑：加载样本、抽状态、训练 shadow 判别。"),
    ("wami/online_gateway.py", "在线拦截接口，用于流式/逐动作安全检查。"),
    ("wami/llm_client.py", "OpenAI-compatible API 客户端，用于接 Qwen 等 base_url/api_key。"),
    ("wami/paper_calibration.py", "阈值/operating point 校准，平衡 IR 与 FPR。"),
    ("wami/adapters.py", "灵活数据读取器，兼容字段名不同的 JSON/JSONL。"),
    ("wami/evaluate.py", "统一指标计算：IR、FPR、ACC、Latency。"),
]


SCRIPT_GROUPS = [
    (
        "数据集转换与自生成训练数据",
        [
            ("scripts/convert_datasets.py", "把 InjecAgent/BIPIA/AgentDojo 官方数据转成 *_wami.jsonl。"),
            ("scripts/generate_self_training_data.py", "生成额外自训练/自监督样本，补充 shadow training。"),
            ("scripts/convert_cyberseceval3_vpi_to_wami.py", "把视觉提示注入数据转成 WAMI 多模态格式。"),
        ],
    ),
    (
        "WAMI 主方法与训练",
        [
            ("scripts/train_wami_paper_strict.py", "训练更接近论文结构的 WAMI torch 模型。"),
            ("scripts/run_paper_mine_gateway.py", "运行最终 paper-faithful WAMI gateway。"),
            ("scripts/run_paper_mine_ensemble.py", "多模型/多种子 ensemble 版本。"),
            ("scripts/run_paper_mine_ensemble_gate.py", "ensemble gate 版本，做更稳健的决策融合。"),
            ("scripts/run_paper_mine_dual_ensemble.py", "双 ensemble 对比版本。"),
            ("scripts/run_paper_experiment_matrix.py", "批量跑论文实验矩阵。"),
            ("scripts/sweep_paper_mine_gateway_configs.py", "扫阈值和配置，找 IR/FPR 合适工作点。"),
        ],
    ),
    (
        "真实 live-agent / Qwen2.5 动作级实验",
        [
            ("scripts/run_qwen_full_live_wami_runtime.py", "本地 qwen2.5:7b-instruct 作为 agent planner，WAMI 在 action 执行前拦截。"),
            ("scripts/run_full_live_wami_runtime.py", "通用 full live runtime。"),
            ("scripts/run_live_planner_next_action.py", "只让模型生成下一步 action 的小规模验证。"),
            ("scripts/recompute_live_wami_action_metrics.py", "把 live-agent 结果重算成 Dangerous Action Generation 和 WAMI Action Block Rate。"),
            ("scripts/trace_wami_decision.py", "输出单条样本的完整 WAMI 判断过程和中间分数。"),
        ],
    ),
    (
        "Table 1 防御方法对比",
        [
            ("scripts/build_final_table1_reproduction.py", "汇总 Table 1 最终认可结果。"),
            ("scripts/run_bookagent_constraint_verifier.py", "复现/近似 BookAgent 安全约束 verifier。"),
            ("scripts/run_webagentguard_paper_method.py", "WebAgentGuard paper-method/action fidelity 版本。"),
            ("scripts/run_webagentguard_noapi_method.py", "无需 API 的 WebAgentGuard 替代流程。"),
            ("scripts/run_guardreasoner_vl_table1.py", "GuardReasoner-VL 本地/量化测试。"),
            ("scripts/run_llamaguard3_ollama_on_datasets.py", "Llama-Guard 3 8B 本地 Ollama 测试。"),
            ("scripts/run_agentdojo_official_table1.py", "AgentDojo 官方检测/spotlighting 相关表一实验。"),
            ("scripts/run_agentdojo_official_detector_on_wami_datasets.py", "AgentDojo 检测器跑 WAMI 三数据集。"),
        ],
    ),
    (
        "Table 2 检测/扰动类 baseline",
        [
            ("scripts/run_table2_official_erase_check.py", "Erase-and-Check 官方/本地模型版本。"),
            ("scripts/run_smoothllm_qwen_judge_on_datasets.py", "SmoothLLM-style + Qwen judge。"),
            ("scripts/run_smoothllm_on_wami_datasets.py", "SmoothLLM-style 本地/非 API 版本。"),
            ("scripts/run_toolemu_sandbox_table2.py", "ToolEmu-Sandbox 表二实验。"),
            ("scripts/run_toolemu_evaluator_on_wami_datasets.py", "ToolEmu evaluator 跑三数据集。"),
            ("scripts/run_official_erase_check.py", "Erase-and-Check 原始/官方入口尝试。"),
            ("scripts/run_official_smoothllm.py", "SmoothLLM 官方入口尝试。"),
            ("scripts/run_official_toolemu.py", "ToolEmu 官方入口尝试。"),
        ],
    ),
    (
        "Table 3 跨模型/跨 agent 实验",
        [
            ("scripts/evaluate_llm_agent.py", "评估不同本地 LLM agent 的动作生成/安全结果。"),
            ("scripts/run_llm_agent.py", "普通 LLM agent 运行入口。"),
            ("scripts/run_wami_runtime_replay.py", "回放 runtime 轨迹并测试 WAMI。"),
        ],
    ),
    (
        "Table 4 ToolBench/AgentBench 能力保持",
        [
            ("scripts/run_toolbench_default_evalset_qwen25_table4.py", "ToolBench 默认 evalset + qwen2.5 表四实验。"),
            ("scripts/run_toolbench_table4_nonlite_baselines.py", "ToolBench 非 lite baseline。"),
            ("scripts/run_toolbench_table4_all_methods.py", "ToolBench 多方法统一表四。"),
            ("scripts/run_toolbench_qwen25_wami_capability.py", "WAMI 对 ToolBench 工具能力保持影响。"),
            ("scripts/run_agentbench_proxy_table4_nonlite_baselines.py", "AgentBench proxy 非 lite baseline。"),
            ("scripts/run_table4_capability_proxy.py", "表四能力保持 proxy 汇总。"),
        ],
    ),
    (
        "Table 5 消融实验",
        [
            ("scripts/run_wami_paper_strict_ablation.py", "最终论文式消融：去掉不同模块看 IR/FPR 变化。"),
            ("scripts/run_wami_paper_ablation.py", "较早 WAMI 消融。"),
            ("scripts/run_wami_ablation_suite.py", "批量消融套件。"),
            ("scripts/run_wami_v4_ablation.py", "v4 结构消融。"),
            ("scripts/build_final_table5_ablation.py", "汇总最终 Table 5。"),
            ("scripts/analyze_wami_v4_decision_contributions.py", "分析各模块对最终决策贡献。"),
        ],
    ),
    (
        "多模态/视觉注入实验",
        [
            ("scripts/run_multimodal_qwenvl_batch.py", "Qwen-VL 多模态批量测试。"),
            ("scripts/run_smoothvlm_style_vpi.py", "SmoothVLM-style 视觉提示注入测试。"),
            ("scripts/run_paper_multimodal_mine.py", "WAMI 多模态 MINE 流程。"),
            ("scripts/run_cyberseceval3_vpi_wami_qwenvl.py", "CyberSecEval3 VPI + Qwen-VL。"),
            ("scripts/run_cyberseceval3_vpi_wami_backend_ablation.py", "多模态 backend 消融。"),
        ],
    ),
    (
        "结果导出、图片和 Word",
        [
            ("scripts/export_method_audit_excels.py", "按方法导出每条样本 Excel：拦截、漏拦、误拦全部保留。"),
            ("scripts/export_wami_accepted_audit_csv.py", "导出最终认可 WAMI paper-faithful 每条样本 CSV。"),
            ("scripts/export_paper_mine_scores.py", "导出 MINE 分数、曲线和 audit 文件。"),
            ("scripts/build_final_experiment_word.py", "生成最终实验结果 Word。"),
            ("scripts/insert_live_wami_rows_into_word_main_tables.py", "把 live-agent action-level 指标插入主表。"),
            ("scripts/plot_final_figures_paper_style.py", "按论文风格画最终图。"),
            ("scripts/plot_final_figures_3_to_7.py", "生成 Figure 3-7。"),
            ("scripts/build_figure7_resource_comparison.py", "生成资源开销图。"),
        ],
    ),
]


FINAL_OUTPUTS = [
    ("data/WAMI最终实验结果汇总.docx", "最终论文实验表格和图片汇总。"),
    ("data/WAMI代码实验全量讲解_导师版.docx", "本文件：全量代码和实验讲解。"),
    ("data/method_audit_excels_expanded/WAMI_paper_faithful.xlsx", "WAMI paper-faithful 每条数据：拦截、漏拦、误拦、分数。"),
    ("data/method_audit_excels_expanded/WAMI_live_agent_action_level_qwen25.xlsx", "qwen2.5 live-agent action-level 每条数据。"),
    ("data/method_audit_excels_expanded/WebAgentGuard_action_fidelity.xlsx", "WebAgentGuard 认可版本每条数据。"),
    ("data/method_audit_excels_expanded/Erase_and_Check_balanced.xlsx", "Erase-and-Check 认可版本每条数据。"),
    ("data/method_audit_excels_expanded/SmoothLLM_style_balanced.xlsx", "SmoothLLM 认可版本每条数据。"),
    ("data/method_audit_excels_expanded/ToolEmu_Sandbox_tau7.xlsx", "ToolEmu-Sandbox 认可版本每条数据。"),
    ("data/method_audit_excels_expanded/Llama_Guard_3_local.xlsx", "Llama-Guard 3 local 认可版本每条数据。"),
    ("data/qwen25_live_wami_recomputed_action_metrics.csv", "live-agent action-level 重算指标。"),
    ("data/wami_paper_faithful_accepted_audit_details.csv", "WAMI paper-faithful 最终认可 per-sample 明细。"),
]


def line_count(path: Path) -> int:
    try:
        return len(path.read_text(encoding="utf-8", errors="ignore").splitlines())
    except Exception:
        return 0


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.7)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Subtitle", 10.5, "44546A"),
        ("Heading 1", 14.5, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
        ("Heading 3", 10.6, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text: str, bold: bool = False, center: bool = False, size: float = 7.6) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)


def table(doc: Document, headers: list[str], rows: list[list[str]], font_size: float = 7.4) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, center=True, size=7.6)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=font_size)
    doc.add_paragraph()


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.14)
    para.paragraph_format.space_after = Pt(3)
    r = para.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(7.3)


def callout(doc: Document, title: str, body: str, fill: str = "EAF4FF") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    shade(c, fill)
    c.text = ""
    r = c.paragraphs[0].add_run(title)
    r.bold = True
    r.font.size = Pt(9)
    p2 = c.add_paragraph()
    p2.add_run(body).font.size = Pt(8.6)
    doc.add_paragraph()


def csv_preview(path: Path, max_rows: int = 5) -> list[list[str]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            rows = []
            for i, row in enumerate(reader):
                rows.append([x[:80] for x in row[:8]])
                if i >= max_rows:
                    break
            return rows
    except Exception:
        return []


def build() -> None:
    doc = Document()
    setup(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("WAMI 全量代码与实验讲解")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("不是只讲核心方法，而是把实际生成的代码、实验脚本、结果文件和论文表格一一对应")

    wami_py = list((ROOT / "wami").glob("*.py"))
    script_py = list((ROOT / "scripts").glob("*.py"))
    wami_lines = sum(line_count(x) for x in wami_py)
    script_lines = sum(line_count(x) for x in script_py)
    xlsx_count = len(list((ROOT / "data" / "method_audit_excels_expanded").glob("*.xlsx")))
    csv_count = len(list((ROOT / "data").glob("*.csv")))

    callout(
        doc,
        "先说结论",
        f"这不是几个脚本的小复现，而是一套完整实验工程：wami 主包 {len(wami_py)} 个 Python 文件，约 {wami_lines} 行；scripts 实验脚本 {len(script_py)} 个，约 {script_lines} 行；"
        f"data 下至少 {csv_count} 个 CSV 结果文件，并且按方法导出了 {xlsx_count} 个 Excel 审计文件。"
        "所以给导师讲时不能只讲 TDG/MINE，还要讲数据转换、训练、baseline、live-agent、消融、表格、图、每条样本审计。",
        "FFF4E5",
    )

    doc.add_heading("1. 这套代码到底复现了论文哪些部分", level=1)
    table(
        doc,
        ["论文部分", "代码实现位置", "已经做到什么程度", "导师听得懂的解释"],
        [
            ["数据集与任务构造", "scripts/convert_datasets.py, wami/datasets.py", "三数据集统一成 intent/plan/label", "把三套格式不同的数据变成同一个 agent 轨迹输入。"],
            ["TDG 工具依赖图", "wami/tdg.py, wami/runtime_trace.py", "静态 plan 和 live runtime 都能建图", "看不可信内容如何跨工具传播。"],
            ["World Model", "wami/torch_model.py", "Transformer + GRU memory + slot memory", "预测 agent 下一步状态变化，不是直接关键词判断。"],
            ["MINE 风险评分", "wami/paper_mine_gateway.py, wami/torch_training.py", "MINE/风险头 + 阈值决策", "给每条轨迹算风险分，超过阈值拦截。"],
            ["Shadow Training", "wami/shadow.py, wami/shadow_llm.py, generate_self_training_data.py", "合成和 LLM 生成攻击变体", "让模型见到更多攻击形态，减少只记住数据集。"],
            ["多模态输入", "wami/multimodal.py, run_smoothvlm_style_vpi.py", "Qwen-VL/本地视觉模型接口和 VPI 测试", "把图片里的提示注入也转成 latent。"],
            ["真实 agent runtime", "wami/live_agent.py, run_qwen_full_live_wami_runtime.py", "qwen2.5 本地模型生成 action，WAMI 动作级拦截", "更接近真实 agent：先生成动作，再判断能不能执行。"],
            ["对比方法", "Table1/Table2/Table3/Table4 相关脚本", "WebAgentGuard、Erase-and-Check、SmoothLLM、ToolEmu、Llama-Guard 等", "不是只跑自己方法，也跑了多个 baseline。"],
            ["消融实验", "run_wami_paper_strict_ablation.py 等", "去掉模块看指标变化", "证明 TDG/MINE/world model/shadow training 各自作用。"],
            ["结果可追溯", "export_method_audit_excels.py", "每个方法一个 Excel，每条样本记录判断", "导师要看具体拦了哪条，也能拿出来。"],
        ],
    )

    doc.add_heading("2. 主包 wami：每个核心文件到底干什么", level=1)
    rows = []
    for rel, desc in CORE_FILES:
        path = ROOT / rel
        rows.append([rel, str(line_count(path)), desc])
    table(doc, ["文件", "行数", "作用"], rows)

    doc.add_heading("3. 主流程按代码走一遍", level=1)
    p(doc, "这一节是给完全不懂代码的人讲的：一条样本从数据集进来，到最后被 WAMI 拦截，中间经过哪些文件。")
    table(
        doc,
        ["顺序", "经过的代码", "输入", "输出", "解释"],
        [
            ["1", "wami/datasets.py", "官方原始数据", "PlanSample(intent, plan, label)", "把不同数据集变成统一样本。"],
            ["2", "wami/tdg.py::build_tdg", "plan", "TDG(nodes, edges)", "把工具调用文本变成图。"],
            ["3", "wami/tdg.py::analyze_tdg_security", "TDG", "source/sink/sensitive/path", "找不可信来源和危险流向。"],
            ["4", "wami/torch_model.py", "intent + TDG 节点序列", "latent state / memory", "预测动作后的安全状态。"],
            ["5", "wami/paper_mine_gateway.py", "latent + security features", "MINE score", "把风险变成数值。"],
            ["6", "threshold compare", "MINE score", "block/allow", "分数超过阈值拦截。"],
            ["7", "wami/evaluate.py", "block/allow + label", "IR/FPR/ACC", "和真实标签对比算实验指标。"],
        ],
    )
    code(
        doc,
        "sample = PlanSample(intent='用户请求', plan='Action: Tool(...)', label=1)\n"
        "tdg = build_tdg(sample.plan)\n"
        "analysis = analyze_tdg_security(tdg)\n"
        "decision = PaperMINEGateway(...).inspect(sample.intent, sample.plan)\n"
        "metrics = evaluate_gateway(gateway, samples)"
    )

    doc.add_heading("4. scripts：实验脚本不是一个，是按论文表格分组的一套", level=1)
    for group, items in SCRIPT_GROUPS:
        doc.add_heading(group, level=2)
        rows = []
        for rel, desc in items:
            rows.append([rel, str(line_count(ROOT / rel)), desc])
        table(doc, ["脚本", "行数", "作用"], rows)

    doc.add_heading("5. 论文表格和代码/结果文件对应关系", level=1)
    table(
        doc,
        ["论文位置", "对应代码", "对应结果文件", "讲解重点"],
        [
            ["Table 1 防御方法对比", "build_final_table1_reproduction.py + 多个 run_* baseline", "data/final_table1_reproduction.csv, data/method_audit_excels_expanded/*.xlsx", "比较 WAMI 和 WebAgentGuard、BookAgent、Llama-Guard 等 IR/FPR/ACC。"],
            ["Table 2 检测/扰动方法", "run_table2_official_erase_check.py, run_smoothllm_*, run_toolemu_*", "data/final_table2_reproduction.csv", "展示 Erase-and-Check、SmoothLLM、ToolEmu 与 WAMI 的差异。"],
            ["Table 3 跨模型/跨 Agent", "evaluate_llm_agent.py, run_llm_agent.py", "data/final_table3_cross_agent_reproduction.csv", "看不同本地模型/agent 设置下结果变化。"],
            ["Table 4 能力保持", "run_toolbench_*, run_agentbench_*", "data/final_table4_capability_reproduction.csv", "不是只看安全，也看正常任务能力有没有被防御破坏。"],
            ["Table 5 消融", "run_wami_paper_strict_ablation.py, build_final_table5_ablation.py", "data/final_table5_ablation.csv", "去掉不同模块后结果变化，说明模块贡献。"],
            ["Figure 3-8", "plot_final_figures_paper_style.py, build_figure7_resource_comparison.py", "data/final_figure*.png/pdf/csv", "把最终认可结果画成论文图。"],
            ["每条样本审计", "export_method_audit_excels.py", "data/method_audit_excels_expanded/*.xlsx", "导师质疑某个数字时，可以打开每条数据核对。"],
        ],
    )

    doc.add_heading("6. 最终/认可结果文件池", level=1)
    rows = []
    for rel, desc in FINAL_OUTPUTS:
        path = ROOT / rel
        rows.append([rel, "存在" if path.exists() else "缺失", str(path.stat().st_size if path.exists() else ""), desc])
    table(doc, ["文件", "状态", "大小(bytes)", "用途"], rows)

    doc.add_heading("7. 三类 WAMI 结果要分清楚", level=1)
    table(
        doc,
        ["版本", "代表文件/脚本", "输入是什么", "指标怎么理解", "论文里怎么写更稳"],
        [
            ["WAMI paper-faithful", "run_paper_mine_gateway.py, WAMI_paper_faithful.xlsx", "数据集里已有 intent+plan", "IR=攻击 plan 被拦截比例；FPR=正常 plan 被误拦比例", "主方法静态轨迹评测。"],
            ["WAMI live-agent action-level", "run_qwen_full_live_wami_runtime.py, recompute_live_wami_action_metrics.py", "intent+observation 给 qwen2.5，生成 action 后 WAMI 拦截", "应看 WAMI Action Block Rate，而不是只看 Old Sample IR", "真实 agent runtime 补充实验。"],
            ["WAMI 多模态", "run_paper_multimodal_mine.py, run_smoothvlm_style_vpi.py", "文本+图片/视觉注入", "看视觉注入是否被拦截", "多模态扩展实验，样本规模要单独说明。"],
        ],
    )

    doc.add_heading("8. 每条数据审计 Excel 是什么", level=1)
    p(doc, "你后来要求“每个方法一个 Excel，里面要有数据集具体数据、拦截/没拦截/误拦”。这个由 export_method_audit_excels.py 完成。")
    audit_dir = ROOT / "data" / "method_audit_excels_expanded"
    rows = []
    for x in sorted(audit_dir.glob("*.xlsx")):
        rows.append([x.name, str(x.stat().st_size), "每条样本的输入、标签、方法输出、是否拦截、是否漏拦/误拦"])
    table(doc, ["Excel", "大小(bytes)", "内容"], rows, font_size=7.0)

    doc.add_heading("9. 当前几个关键 CSV 的表头预览", level=1)
    for rel in [
        "data/final_table1_reproduction.csv",
        "data/final_table2_reproduction.csv",
        "data/final_table3_cross_agent_reproduction.csv",
        "data/final_table4_capability_reproduction.csv",
        "data/final_table5_ablation.csv",
        "data/qwen25_live_wami_recomputed_action_metrics.csv",
    ]:
        path = ROOT / rel
        doc.add_heading(rel, level=2)
        rows = csv_preview(path, 3)
        if rows:
            table(doc, [f"列{i+1}" for i in range(max(len(r) for r in rows))], rows, font_size=6.8)
        else:
            p(doc, "未找到或无法读取。")

    doc.add_heading("10. 给导师讲代码时的重点话术", level=1)
    bullet(doc, "第一句先讲工程规模：主包 27 个文件，实验脚本 128 个，覆盖数据转换、模型训练、baseline、live-agent、消融、图表和样本审计。")
    bullet(doc, "第二句讲核心创新：WAMI 不是问大模型“是否攻击”，而是在 agent 执行动作前，把工具调用变成 TDG，再用 world model 和 MINE 分数判断风险。")
    bullet(doc, "第三句讲数据：三数据集原始格式不同，adapter 只是统一成 agent trajectory，label 不输入模型。")
    bullet(doc, "第四句讲 Qwen：Qwen 在 live-agent 里只是生成下一步 action，WAMI 才是防御器。")
    bullet(doc, "第五句讲可信度：每个方法都有 per-sample Excel，能追到具体哪条数据被拦、哪条漏拦、哪条误拦。")

    doc.add_heading("11. 这版相比上一版补强了什么", level=1)
    table(
        doc,
        ["上一版问题", "这版补强"],
        [
            ["只讲主链路，没有体现代码量", "加入主包/脚本数量、行数、核心文件列表。"],
            ["实验脚本讲得少", "按 Table1-5、Figure、多模态、live-agent、导出审计分组讲。"],
            ["导师不知道结果从哪个脚本来", "加入表格/脚本/结果文件三方对应。"],
            ["每条样本审计没强调", "单独列出 method_audit_excels_expanded 的 Excel。"],
            ["没有体现你做过的争议实验和认可结果池", "把最终结果文件、CSV、Excel、Word 都列出来。"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

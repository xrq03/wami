from __future__ import annotations

from pathlib import Path
from textwrap import shorten

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "WAMI逐文件代码讲解手册.docx"


FILES = [
    {
        "path": "wami/tdg.py",
        "role": "把 agent 的文本 plan 解析成工具依赖图 TDG。",
        "why": "如果不先把 plan 变成图，WAMI 就不知道外部 observation 是怎么一步步流向 GmailSendEmail、ExfiltrateData 这类危险工具的。",
        "functions": [
            ("build_tdg", "主入口。输入一段 plan 文本，输出 TDG(nodes, edges)。"),
            ("_parse_args", "把工具调用括号里的参数解析成 dict。"),
            ("topological_order", "按照依赖顺序遍历工具节点。"),
        ],
        "talk": "这个文件就是 WAMI 的“读图器”。它把 Action: Tool(...) 这种文本工具调用变成图结构，后面的 world model 和 gateway 都依赖这个图。",
    },
    {
        "path": "wami/torch_model.py",
        "role": "实现 WAMI 的神经网络版本，包括 world model、memory、slot memory、MINE 打分头。",
        "why": "这是 WAMI 的模型主体，用来预测每一步工具动作后的状态，并给出意图-状态对齐分数。",
        "functions": [
            ("TorchWAMIConfig.paper_strict", "论文式模型配置：1024 维、4 层 Transformer、8 heads。"),
            ("next_state", "world model 的核心，输入当前状态和动作，预测下一状态。"),
            ("score", "MINE 主打分函数，判断 intent 和 state 是否对齐。"),
            ("transition_score", "辅助打分，判断状态转移是否异常。"),
        ],
        "talk": "这个文件是 WAMI 的“大脑”。它不是调用大模型判断安全，而是本地神经网络根据工具动作预测状态，再用 MINE 分数判断偏离程度。",
    },
    {
        "path": "wami/paper_mine_gateway.py",
        "role": "实现最终 paper-faithful WAMI gateway。",
        "why": "真正做 allow/block 决策的地方。它遍历 TDG 中每一步动作，低于阈值就拦截。",
        "functions": [
            ("PaperMINEConfig", "保存阈值、risk_margin、transition_fusion 等配置。"),
            ("PaperMINEGateway.inspect", "主入口。输入 intent 和 plan，输出 GateDecision。"),
            ("_risk_margin", "根据工具风险调整阈值，但不是直接强规则拦截。"),
        ],
        "talk": "这个文件就是“安检口”。agent 每准备执行一个动作，gateway 就用 MINE 分数检查是否还符合用户原始目标。",
    },
    {
        "path": "scripts/export_wami_accepted_audit_csv.py",
        "role": "重新导出最终认可的 WAMI paper-faithful 逐样本结果。",
        "why": "Word 主表里的 WAMI 结果要和 Excel 逐样本审计对齐，所以这个脚本按最终采用口径重新跑一遍每条样本。",
        "functions": [
            ("main", "加载 source-aware 和 triplet-slot 两个 checkpoint，逐条样本判断。"),
            ("make_gateway", "按 tau 和 fusion 参数创建 PaperMINEGateway。"),
            ("write_csv", "保存逐样本 CSV。"),
        ],
        "talk": "这个脚本的作用是把最终 WAMI 主表背后的每一条样本判断都保存下来，方便导师追溯。",
    },
    {
        "path": "scripts/recompute_live_wami_action_metrics.py",
        "role": "重新计算 qwen2.5 live-agent action-level 指标。",
        "why": "旧 sample-level IR 会把 qwen2.5 没生成危险动作也算进攻击总数，低估 WAMI。这个脚本专门按危险动作是否进入 WAMI 来重算。",
        "functions": [
            ("main", "读取 qwen25_7b_ollama_full_live_wami_*_full.csv，计算动作级指标。"),
            ("rate", "安全地计算比例。"),
            ("to_md", "把结果写成 Markdown 表格。"),
        ],
        "talk": "这个脚本解决一个关键问题：WAMI 只能拦截已经生成的危险动作，不能因为 agent 没生成危险动作就算 WAMI 漏拦。",
    },
    {
        "path": "scripts/export_method_audit_excels.py",
        "role": "把所有方法的结果统一导出成每个方法一个 Excel。",
        "why": "导师要看具体数据，所以不能只给 IR/FPR，需要每条样本的原始请求、plan、observation、注入、危险动作和 TP/FN/FP/TN。",
        "functions": [
            ("main", "组织所有方法的导出任务。"),
            ("build_wami", "生成 WAMI paper-faithful Excel。"),
            ("build_live_wami_action_level", "生成 qwen2.5 live-agent 动作级 Excel。"),
            ("extract_concrete_fields", "从样本中抽取原始请求、外部 observation、注入指令、危险动作。"),
            ("write_method_workbook", "写 Excel，并分 Summary、每条具体数据、TP/FN/FP/TN 等 sheet。"),
        ],
        "talk": "这个文件不是算法本身，而是证据整理器。它把每个方法的每条判断结果都整理成导师能打开看的 Excel。",
    },
    {
        "path": "scripts/build_final_experiment_word.py",
        "role": "生成最终实验汇总 Word。",
        "why": "把所有表格、图、最终认可结果、备选结果和运行代码映射整理到一个 Word。",
        "functions": [
            ("main", "创建 Word，写入表格、图、最终解释。"),
            ("add_table", "往 Word 中写表格。"),
            ("add_figure", "往 Word 中插入实验图。"),
        ],
        "talk": "这个脚本负责生成导师看到的总报告，不负责模型判断。",
    },
    {
        "path": "scripts/insert_live_wami_rows_into_word_main_tables.py",
        "role": "把 live-agent action-level 结果插入 Word 主表。",
        "why": "用户要求 live-agent action-level 结果不要只放在最后说明，而是直接进入主表。",
        "functions": [
            ("add_to_main_table", "把 live-agent 三个数据集结果加到 Table A。"),
            ("add_to_macro_table", "计算并加入 macro 平均。"),
            ("add_to_defense_table", "加到 Table 1。"),
            ("add_to_frontier_table", "加到 Table 2。"),
        ],
        "talk": "这个脚本是 Word 后处理，把动作级 WAMI 结果也放进主表里。",
    },
    {
        "path": "scripts/run_bookagent_constraint_verifier.py",
        "role": "BookAgent-style baseline。",
        "why": "复现 BookAgent 的安全约束思路：先从用户 intent 提取允许行为，再看 plan 是否越权。",
        "functions": [
            ("inspect", "检查一条 intent+plan 是否违反约束。"),
            ("extract_constraints", "从用户请求抽取允许动作和敏感项。"),
            ("evaluate", "对一个数据集计算 IR/FPR/ACC。"),
        ],
        "talk": "这个 baseline 看的是 agent 是否超出用户授权范围，比如用户只是查询，却出现发邮件或泄露数据。",
    },
    {
        "path": "scripts/run_table2_official_erase_check.py",
        "role": "Erase-and-Check baseline。",
        "why": "模拟擦除输入部分内容后再检查是否危险的防御方法。",
        "functions": [
            ("_evaluate", "逐样本调用 erase_and_check，统计指标。"),
            ("load_raw_injecagent", "加载 InjecAgent 原始样本。"),
            ("load_raw_bipia", "加载 BIPIA 原始样本。"),
            ("load_raw_agentdojo", "加载 AgentDojo 原始样本。"),
        ],
        "talk": "这个方法主要作为对比，不是 WAMI 本身。它关注文本擦除后的安全判断。",
    },
    {
        "path": "scripts/run_smoothllm_qwen_judge_on_datasets.py",
        "role": "SmoothLLM-style baseline。",
        "why": "对输入做扰动，然后用 judge 判断是否 dangerous，观察扰动后判断是否稳定。",
        "functions": [
            ("_evaluate", "对每条样本生成扰动副本并投票。"),
            ("_judge_with_cache", "优先用缓存，避免重复消耗 API/token。"),
            ("_ollama_judge", "本地 Ollama 模型判断 harmful/safe。"),
        ],
        "talk": "这个 baseline 的思想是：如果输入稍微扰动后仍然触发危险判断，就说明它可能有问题。",
    },
]


def setup(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8.0)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    doc.add_heading(title, level=2)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    doc.add_paragraph()


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.18)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.2)


def snippet(path_text: str, keywords: list[str], max_lines: int = 22) -> str:
    path = ROOT / path_text
    if not path.exists():
        return "[文件不存在]"
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    hits: list[str] = []
    for keyword in keywords:
        for i, line in enumerate(lines):
            if keyword in line:
                start = max(0, i - 2)
                end = min(len(lines), i + max_lines)
                block = []
                for n in range(start, end):
                    block.append(f"{n + 1:04d}: {lines[n]}")
                hits.append("\n".join(block))
                break
    return "\n\n".join(hits) if hits else "\n".join(f"{i+1:04d}: {line}" for i, line in enumerate(lines[:max_lines]))


def main() -> None:
    doc = Document()
    setup(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("WAMI 逐文件代码讲解手册")
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(11, 37, 69)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("给导师讲代码用：每个文件干什么、入口函数是什么、关键代码怎么读").italic = True

    doc.add_heading("0. 这份手册怎么用", level=1)
    p(doc, "你可以把这份文档当成代码导游。导师问到某个模块时，你不用临时翻代码，可以直接按这里的“文件作用、入口函数、关键代码片段、讲法”说明。")
    p(doc, "最重要的顺序是：先讲 wami 目录里的核心算法，再讲 scripts 目录里的实验脚本，最后讲 data 目录里的输出结果。")

    table(
        doc,
        "代码层级总览",
        ["层级", "代表文件", "一句话解释"],
        [
            ["核心算法层", "wami/tdg.py, wami/torch_model.py, wami/paper_mine_gateway.py", "实现 WAMI 方法本体。"],
            ["实验运行层", "scripts/run_*.py", "跑 WAMI 和各个 baseline。"],
            ["结果整理层", "scripts/build_final_experiment_word.py, scripts/export_method_audit_excels.py", "生成 Word 和 Excel。"],
            ["证据输出层", "data/*.docx, data/method_audit_excels_expanded/*.xlsx", "给导师看的最终结果和逐样本证据。"],
        ],
    )

    for item in FILES:
        doc.add_heading(item["path"], level=1)
        p(doc, f"这个文件干什么：{item['role']}")
        p(doc, f"为什么需要它：{item['why']}")
        p(doc, f"给导师怎么讲：{item['talk']}")
        table(
            doc,
            "这个文件里的关键函数",
            ["函数/类", "作用"],
            [[name, desc] for name, desc in item["functions"]],
        )
        keywords = [name.split(".")[-1] for name, _ in item["functions"][:3]]
        doc.add_heading("关键代码片段应该怎么看", level=2)
        code(doc, snippet(item["path"], keywords))
        doc.add_heading("你可以这样讲这段代码", level=2)
        for name, desc in item["functions"]:
            p(doc, f"- {name}：{desc}")
        doc.add_page_break()

    doc.add_heading("最后：给导师讲代码的标准路线", level=1)
    p(doc, "第一步，打开 wami/tdg.py：说明我先把工具调用 plan 变成图。")
    p(doc, "第二步，打开 wami/torch_model.py：说明我用本地神经网络预测每一步工具动作后的状态，并用 MINE 输出对齐分数。")
    p(doc, "第三步，打开 wami/paper_mine_gateway.py：说明真正的 allow/block 在这里发生，低于动态阈值就拦截。")
    p(doc, "第四步，打开 scripts/recompute_live_wami_action_metrics.py：说明为什么 live-agent 结果要按危险动作拦截率重算。")
    p(doc, "第五步，打开 scripts/export_method_audit_excels.py：说明每个方法都有逐样本 Excel，导师可以检查具体数据。")
    p(doc, "最后打开 data/WAMI最终实验结果汇总.docx 和 data/method_audit_excels_expanded 里的 Excel，证明代码结果已经整理成论文表格和逐样本证据。")

    table(
        doc,
        "如果导师问“你这个代码到底跑了什么”",
        ["回答点", "你可以直接说"],
        [
            ["输入", "BIPIA、InjecAgent、AgentDojo 数据集，每条样本包括用户请求、工具调用 plan、标签。"],
            ["方法", "WAMI 先构建 TDG，再用 world model rollout，最后用 MINE/Gateway 判断拦截。"],
            ["输出", "IR、FPR、ACC、延迟、动作级拦截率，以及每条样本的 TP/FN/FP/TN Excel。"],
            ["可复查性", "所有结果都在 data 目录，所有生成脚本都在 scripts 目录，导师可以从 Excel 看到具体样本。"],
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

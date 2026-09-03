from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "data" / "WAMI最终实验结果汇总.docx"
LIVE_CSV = ROOT / "data" / "qwen25_live_wami_recomputed_action_metrics.csv"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def pct(value: float) -> str:
    return f"{float(value) * 100:.1f}%"


def main() -> None:
    doc = Document(DOCX)
    df = pd.read_csv(LIVE_CSV)

    doc.add_page_break()
    doc.add_heading("13. WAMI 实验流程口径标注与 live-agent 重算结果", level=1)

    doc.add_paragraph(
        "本节用于区分 Word 中已有 WAMI 主结果和 qwen2.5 live-agent 运行结果，避免把 agent 是否生成危险动作与 WAMI 是否拦截危险动作混在一起。"
    )

    doc.add_heading("13.1 Word 已有 WAMI 主表结果的流程", level=2)
    doc.add_paragraph(
        "Word 主表中的 WAMI paper-faithful 结果采用的是数据集原始 plan / tool trajectory 评测流程："
        "数据集已经给出用户请求、工具调用顺序、外部 observation、注入指令以及最终危险动作，WAMI 直接对该工具轨迹进行 TDG 构建、world model rollout 和 MINE/Gateway 判断。"
        "因此该结果衡量的是：当攻击轨迹或危险动作已经出现在待审计 plan 中时，WAMI 能不能识别并拦截。"
    )
    doc.add_paragraph(
        "该流程对应 Word 中 Table A、Table 1、Table 2 和 Table 5 的 WAMI paper-faithful 行。"
        "最终采用口径为：BIPIA 与 InjecAgent 使用 source-aware + triplet-slot ensemble OR；AgentDojo 使用 triplet-slot tau50，以降低 AgentDojo 上的误报率。"
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Dataset", "Flow", "IR", "FPR", "ACC"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in [
        ["BIPIA", "paper-faithful dataset-plan audit", "99.8%", "0.5%", "99.6%"],
        ["InjecAgent", "paper-faithful dataset-plan audit", "86.8%", "5.9%", "90.5%"],
        ["AgentDojo", "paper-faithful dataset-plan audit", "97.2%", "9.3%", "96.3%"],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    doc.add_heading("13.2 qwen2.5 live-agent action-level 重算流程", level=2)
    doc.add_paragraph(
        "qwen25_7b_ollama_full_live_wami_*_full.csv 对应的是另一套流程：先由本地 qwen2.5:7b-instruct 作为 live agent 根据用户请求和 observation 生成动作，"
        "然后 WAMI gateway 在动作执行前判断是否拦截。这个流程会受到 agent 是否真的生成危险动作影响。"
    )
    doc.add_paragraph(
        "因此这里不把“agent 没生成危险动作”算成 WAMI 漏拦，而是拆成两个指标："
        "Dangerous Action Generation Rate 表示攻击样本中 agent 实际生成危险动作的比例；"
        "WAMI Action Block Rate 表示在这些已经生成的危险动作里，WAMI 实际拦截了多少。"
        "误拦截也同时给出 sample-level FPR 和 action-level false block rate。"
    )

    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = [
        "Dataset",
        "Attack N",
        "Old Sample IR",
        "Dangerous Action Gen.",
        "Dangerous Action N",
        "WAMI Action Block",
        "Released Dangerous",
        "Benign Sample FPR",
        "Benign Action False Block",
        "Latency ms",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for _, row in df.iterrows():
        values = [
            row["dataset"],
            int(row["attack_n"]),
            pct(row["old_sample_ir"]),
            pct(row["dangerous_action_generation_rate"]),
            int(row["dangerous_action_n"]),
            pct(row["wami_action_block_rate"]),
            int(row["dangerous_action_released_n"]),
            pct(row["fpr_on_benign_samples"]),
            pct(row["benign_action_false_block_rate"]),
            f"{float(row['latency_ms']):.1f}",
        ]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)

    doc.add_paragraph(
        "解释：Old Sample IR = 攻击样本被拦截数 / 全部攻击样本数。该指标在 live-agent 流程中会低估 WAMI，"
        "因为部分攻击样本中 qwen2.5 并没有走到危险工具动作。WAMI Action Block Rate = 已拦截危险动作 / agent 实际生成的危险动作，"
        "更适合说明 WAMI gateway 本身对危险动作的拦截能力。"
    )
    doc.add_paragraph(
        "结论：live-agent 结果说明，在 qwen2.5 实际生成危险动作的情况下，WAMI 对 BIPIA、InjecAgent、AgentDojo 的危险动作拦截率分别为 "
        "100.0%、89.3%、93.1%；同时正常动作误拦率分别为 0.1%、0.0%、1.8%。"
    )

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "WAMI代码讲解给导师版.docx"


def setup(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8.3)


def table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    doc.add_heading(title, level=2)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(8.8)


def para(doc: Document, text: str, bold_start: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start) :])
    else:
        p.add_run(text)


def main() -> None:
    doc = Document()
    setup(doc)
    sec = doc.sections[0]
    sec.top_margin = Pt(45)
    sec.bottom_margin = Pt(45)
    sec.left_margin = Pt(45)
    sec.right_margin = Pt(45)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("WAMI 复现代码给导师讲解稿")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(11, 37, 69)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("面向零基础听众：每个实验结果对应哪个代码文件、输入是什么、输出是什么、应该怎么讲").italic = True

    doc.add_heading("0. 你给导师先讲的 1 分钟版本", level=1)
    para(
        doc,
        "一句话：这套代码不是只算一个分数，而是把 WAMI 防御框架、对比方法、live-agent 流程和逐样本审计全部串起来。"
    )
    para(
        doc,
        "我可以这样讲：WAMI 是一个 agent 动作执行前的安全网。agent 先产生工具调用计划或下一步工具动作，WAMI 把这些工具调用解析成图，再用 world model 预测状态变化，最后用 MINE 分数判断这个动作是否还符合用户原始意图。如果不符合，就在工具真正执行前拦截。"
    )
    code(
        doc,
        "用户请求 + 工具返回/observation\n"
        "        ↓\n"
        "agent 产生 plan 或下一步 action\n"
        "        ↓\n"
        "TDG: 把工具调用变成依赖图\n"
        "        ↓\n"
        "World Model: 预测每一步动作后的状态\n"
        "        ↓\n"
        "MINE/Gateway: 打分并决定 allow/block\n"
        "        ↓\n"
        "输出 IR、FPR、ACC、逐样本 TP/FN/FP/TN Excel",
    )

    doc.add_heading("1. 总体文件应该怎么讲", level=1)
    table(
        doc,
        "核心文件地图",
        ["你给导师看的东西", "文件路径", "它证明什么"],
        [
            [
                "最终实验总表和图",
                r"D:\论文111\data\WAMI最终实验结果汇总.docx",
                "论文表格、图、baseline、消融、live-agent action-level 结果都在这里。",
            ],
            [
                "WAMI 主方法逐样本审计",
                r"D:\论文111\data\method_audit_excels_expanded\WAMI_paper_faithful.xlsx",
                "每条样本的用户请求、plan、外部 observation、注入指令、危险动作、是否拦截。",
            ],
            [
                "WAMI live-agent 动作级审计",
                r"D:\论文111\data\method_audit_excels_expanded\WAMI_live_agent_action_level_qwen25.xlsx",
                "qwen2.5 真正生成动作后，WAMI 对危险动作的拦截率和正常动作误拦率。",
            ],
            [
                "所有 baseline 审计表",
                r"D:\论文111\data\method_audit_excels_expanded",
                "每个对比方法一个 Excel，可以追溯到具体哪些样本被拦、漏拦、误拦。",
            ],
            [
                "自动导出审计 Excel 的代码",
                r"D:\论文111\scripts\export_method_audit_excels.py",
                "把各方法结果统一整理成逐样本 Excel。",
            ],
        ],
    )

    doc.add_heading("2. WAMI 主流程怎么结合代码讲", level=1)
    para(
        doc,
        "你可以把 WAMI 分成四个零件讲：TDG、World Model、MINE、Gateway。每个零件都有对应代码，不是口头概念。",
    )
    table(
        doc,
        "WAMI 四个核心模块",
        ["模块", "代码文件", "小白解释", "导师可能关心点"],
        [
            [
                "TDG 工具依赖图",
                r"wami\tdg.py",
                "把一行行 Action 工具调用解析成节点和边。节点是工具，边表示后一步用了前一步输出。",
                "说明 WAMI 不是只看关键词，而是看工具调用之间的数据流。",
            ],
            [
                "World Model",
                r"wami\torch_model.py",
                "像一个模拟器，预测执行某个工具动作之后，agent 内部状态会变成什么样。",
                "这里使用 Transformer Encoder、GRU memory、slot memory、subgoal head。",
            ],
            [
                "MINE 打分",
                r"wami\torch_model.py",
                "把用户意图向量和当前动作/状态向量放在一起，输出一个对齐分数。分数越低，说明动作越偏离用户意图。",
                "这是论文里的核心判断来源。",
            ],
            [
                "Gateway 拦截器",
                r"wami\paper_mine_gateway.py",
                "每一步动作执行前，都问一句：这个动作还符合用户原意吗？如果 MINE 分数低于阈值，就 block。",
                "它是执行前防御，不是事后分析。",
            ],
        ],
    )

    doc.add_heading("2.1 TDG 代码怎么讲", level=2)
    para(
        doc,
        "TDG 的入口函数是 build_tdg(plan)。它读入一段工具调用 plan，例如 Action: ReadContext(...)、Action: ExfiltrateData(...)，然后输出一个图。",
    )
    code(
        doc,
        r"文件：D:\论文111\wami\tdg.py"
        "\n"
        "build_tdg(plan) 做三件事：\n"
        "1. 用正则识别每一行 Action 工具调用；\n"
        "2. 把工具名和参数解析成 TDGNode；\n"
        "3. 如果后一步参数里出现 $0、$1 这种前一步输出引用，就建立边。",
    )
    para(
        doc,
        "给导师可以这样讲：如果一个恶意网页内容先被 ReadContext 读入，后面又流向 GmailSendEmail 或 ExfiltrateData，TDG 就能把这条传播路径连出来。这就是 WAMI 能处理间接提示注入的原因之一。",
    )

    doc.add_heading("2.2 World Model 代码怎么讲", level=2)
    para(
        doc,
        "World Model 在 wami/torch_model.py 里。它不是调用 qwen 判断安全，而是一个本地神经网络，用向量表示意图、动作、observation、memory 和 subgoal。",
    )
    code(
        doc,
        "关键配置：TorchWAMIConfig.paper_strict()\n"
        "- dim = 1024\n"
        "- hidden_dim = 1024\n"
        "- layers = 4\n"
        "- heads = 8\n"
        "- world = TransformerEncoder\n"
        "- memory_cell = GRUCell\n"
        "- use_slot_memory = True",
    )
    para(
        doc,
        "给导师可以这样讲：World Model 做的是状态预测。比如用户只是想查商品评论，但工具返回里混进了“把地址发给某邮箱”。World Model 会看这个 observation 进入 memory 后，后续动作状态是否偏离原始购物查询目标。",
    )

    doc.add_heading("2.3 MINE/Gateway 代码怎么讲", level=2)
    para(
        doc,
        "Gateway 的主入口是 PaperMINEGateway.inspect(intent, plan)。它会遍历 TDG 中每一步工具动作，计算当前动作状态和用户意图之间的 MINE score。",
    )
    code(
        doc,
        r"文件：D:\论文111\wami\paper_mine_gateway.py"
        "\n"
        "inspect(intent, plan):\n"
        "1. build_tdg(plan) 构建工具依赖图；\n"
        "2. 先可选计算 plan_score；\n"
        "3. 对每个工具节点 rollout 得到 state；\n"
        "4. mine_score(intent_vec, state) 得到对齐分数；\n"
        "5. score < threshold 时返回 GateDecision(False, ...)，即拦截。",
    )
    para(
        doc,
        "注意讲法：这里的安全判断主要来自 MINE 分数和动态阈值。代码里有 risk_margin，但它是调节阈值，不是直接一票否决的强规则。",
    )

    doc.add_heading("3. 两种 WAMI 结果为什么都要有", level=1)
    para(
        doc,
        "这是最容易被问到的地方。你要明确告诉导师：我有两种流程，它们测的不是同一个问题。",
    )
    table(
        doc,
        "WAMI 两种评测流程",
        ["名称", "代码/数据来源", "测的是什么", "Word 中怎么写"],
        [
            [
                "WAMI paper-faithful",
                r"paper_mine_ensemble_sourceaware_triplet_taua45_taub50_results.csv；paper_mine_triplet_slot_seed4071_e4_tau50_results.csv",
                "数据集已经给出完整危险 plan，直接测试 WAMI 能不能拦截这个危险工具轨迹。",
                "Table A / Table 1 / Table 2 的 WAMI paper-faithful 行。",
            ],
            [
                "WAMI live-agent action-level (qwen2.5)",
                r"qwen25_7b_ollama_full_live_wami_*_full.csv；recompute_live_wami_action_metrics.py",
                "先让 qwen2.5 作为 agent 生成动作，再看 WAMI 对真正生成的危险动作是否拦截。",
                "主表新增的 WAMI live-agent action-level (qwen2.5) 行。",
            ],
        ],
    )
    para(
        doc,
        "导师如果问：为什么 live-agent 的 Old Sample IR 比较低？你答：因为很多攻击样本里 qwen2.5 没有真正走到危险动作，这不是 WAMI 漏拦。所以我重新计算了 WAMI Action Block Rate，只在危险动作真的出现时统计 WAMI 是否拦截。",
    )
    table(
        doc,
        "live-agent action-level 最终结果",
        ["Dataset", "危险动作数", "WAMI 拦截危险动作", "WAMI Action Block", "正常动作误拦率"],
        [
            ["BIPIA", "992", "992", "100.0%", "0.1%"],
            ["InjecAgent", "921", "822", "89.3%", "0.0%"],
            ["AgentDojo", "130", "121", "93.1%", "1.8%"],
        ],
    )

    doc.add_heading("4. 数据集和每条样本怎么讲", level=1)
    para(
        doc,
        "数据集不是只给一句 prompt。我们最终整理后的 Excel 已经把每条数据展开成了导师能看懂的字段。",
    )
    table(
        doc,
        "Excel 逐样本字段含义",
        ["字段", "解释", "讲给导师的话"],
        [
            ["original_user_request", "用户原始请求", "用户真正想让 agent 做的事。"],
            ["agent_plan", "工具调用轨迹", "agent 准备按哪些工具步骤完成任务。"],
            ["tool_sequence", "工具顺序", "把 plan 简化成工具名链条，方便看数据流。"],
            ["external_observation", "外部 observation", "网页、邮件、评论、工具返回等不可信内容。"],
            ["injected_instruction", "注入指令", "攻击者藏在 observation 里的恶意目标。"],
            ["dangerous_or_sink_action", "危险动作", "如发邮件、泄露数据、转账、删除、外发等动作。"],
            ["audit_category", "TP/FN/FP/TN", "这条数据到底是正确拦截、漏拦、误拦还是正确放行。"],
        ],
    )
    para(
        doc,
        "你可以现场打开 WAMI_paper_faithful.xlsx 的“每条具体数据”sheet，给导师看某一行：原始请求只是查询信息，但 plan 中出现 ParseInjectedInstruction 和 ExfiltrateData，WAMI 把它拦住，所以这条是 TP。",
    )

    doc.add_heading("5. 对比方法代码怎么讲", level=1)
    table(
        doc,
        "Baseline 与代码对应关系",
        ["方法", "运行/整理代码", "给导师的解释口径"],
        [
            [
                "SmoothLLM-style",
                r"scripts\run_smoothllm_qwen_judge_on_datasets.py",
                "对输入做扰动，多次判断是否仍然危险；这里用本地 qwen2.5 judge 和 cache，避免重复消耗 token。",
            ],
            [
                "Erase-and-Check",
                r"scripts\run_table2_official_erase_check.py",
                "擦除部分输入再检查任务是否仍然危险，用来模拟 certified-llm-safety 的 erase-and-check 思路。",
            ],
            [
                "ToolEmu-Sandbox",
                r"scripts\run_toolemu_sandbox_table2.py",
                "把工具轨迹放进 sandbox 风险评分，判断工具链是否会导致外发、泄露或副作用。",
            ],
            [
                "Llama-Guard 3",
                r"scripts\run_llamaguard3_ollama_on_datasets.py",
                "本地 Ollama 跑 llama-guard3:8b，作为通用内容安全 guard baseline。",
            ],
            [
                "WebAgentGuard",
                r"scripts\run_webagentguard_paper_method.py",
                "按 WebAgentGuard 的 action fidelity 思路，检查下一步 action 是否采纳或传播不可信 observation。",
            ],
            [
                "BookAgent-style",
                r"scripts\run_bookagent_constraint_verifier.py",
                "复现 BookAgent 安全约束思想：从用户意图抽取允许动作，再检查工具轨迹是否越权。",
            ],
            [
                "GuardReasoner-VL",
                r"scripts\run_guardreasoner_vl_table1.py",
                "用本地/量化视觉安全推理模型做安全判断，结果作为可复现 baseline。",
            ],
        ],
    )

    doc.add_heading("6. 审计 Excel 是怎么由代码生成的", level=1)
    para(
        doc,
        "所有逐样本 Excel 由 scripts/export_method_audit_excels.py 统一生成。这个脚本的作用是把不同 baseline 的不同输出格式整理成同一种表格格式。",
    )
    code(
        doc,
        r"运行命令："
        "\n"
        r"C:\Users\ruoan\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe .\scripts\export_method_audit_excels.py"
        "\n\n"
        r"输出目录：D:\论文111\data\method_audit_excels_expanded"
    )
    para(
        doc,
        "这个脚本做的事情很适合给导师讲：它不是改实验结果，而是把已有实验日志整理成可审计证据。每个方法一个 Excel，每个 Excel 都有 Summary、每条具体数据、TP/FN/FP/TN 等 sheet。",
    )
    table(
        doc,
        "Excel 生成脚本里的关键函数",
        ["函数名", "作用", "怎么讲"],
        [
            ["build_wami", "生成 WAMI paper-faithful 审计表", "对应论文主表的 WAMI 主方法。"],
            ["build_live_wami_action_level", "生成 qwen2.5 live-agent 动作级审计表", "把危险动作生成、动作拦截、动作误拦拆开。"],
            ["build_smoothllm_balanced", "从 cache 还原 SmoothLLM 逐样本结果", "不重新烧 token，直接复用已有判断日志。"],
            ["build_erase_balanced", "从 cache 还原 Erase-and-Check 逐样本结果", "把每条样本对应到 blocked/allowed。"],
            ["add_data_text_column", "抽取具体数据字段", "把原始请求、observation、注入指令、危险动作展开给人看。"],
            ["write_method_workbook", "写 Excel", "按 Summary、每条具体数据、TP/FN/FP/TN 分 sheet 保存。"],
        ],
    )

    doc.add_heading("7. 你可以照着讲的一段完整话", level=1)
    para(
        doc,
        "老师，我这里代码主要分成三层。第一层是 WAMI 方法本身，在 wami/tdg.py 里把 agent 的工具调用解析成依赖图，在 wami/torch_model.py 里用 Transformer world model 预测每步动作后的状态，在 wami/paper_mine_gateway.py 里用 MINE 分数和动态阈值决定是否拦截。第二层是实验运行脚本，分别跑 WAMI 和 SmoothLLM、Erase-and-Check、ToolEmu、Llama-Guard 3、WebAgentGuard、BookAgent-style 等 baseline。第三层是审计整理脚本 export_method_audit_excels.py，它把每个方法的每条样本都导出成 Excel，能看到原始用户请求、外部 observation、注入指令、危险动作、是否拦截，以及 TP/FN/FP/TN 分类。",
    )
    para(
        doc,
        "我还专门区分了两种 WAMI 评测：paper-faithful 是直接在数据集给定的工具轨迹上测 WAMI 拦截能力；live-agent action-level 是先让 qwen2.5 生成动作，再看 WAMI 对实际生成的危险动作是否拦截。这样不会把“agent 没生成危险动作”误算成 WAMI 漏拦，所以指标更合理。",
    )

    doc.add_heading("8. 导师可能问的问题和回答", level=1)
    table(
        doc,
        "问答准备",
        ["导师可能问", "你可以这样回答"],
        [
            [
                "你的 WAMI 是不是靠关键词规则？",
                "不是。核心判断来自 TDG + world model rollout + MINE score。代码里的 risk margin 只是调节阈值，不是直接强规则拦截。",
            ],
            [
                "qwen2.5 在 WAMI 里起什么作用？",
                "在 live-agent 实验里，qwen2.5 是被保护的 agent，用来生成动作；WAMI 是动作执行前的 gateway，不是让 qwen2.5 判断安全。",
            ],
            [
                "为什么有两个 WAMI 结果？",
                "paper-faithful 测数据集原始危险 plan 的拦截能力；live-agent 测真实 agent 生成动作后的 action-level 拦截能力。两个问题不同，所以都报告。",
            ],
            [
                "结果能不能追溯？",
                "可以。每个方法都有单独 Excel，每条样本都保留原始请求、plan、observation、注入内容、危险动作、blocked 和 TP/FN/FP/TN。",
            ],
            [
                "为什么 live-agent 旧 IR 低？",
                "旧 IR 把 agent 没生成危险动作也算进攻击总数，会低估 WAMI。所以我重新计算 WAMI Action Block Rate，只看危险动作真的出现后 WAMI 是否拦截。",
            ],
        ],
    )

    doc.add_heading("9. 最后建议你给导师发哪些文件", level=1)
    table(
        doc,
        "推荐发送文件",
        ["优先级", "文件", "用途"],
        [
            ["1", r"D:\论文111\data\WAMI最终实验结果汇总.docx", "总结果、表格、图、流程标注。"],
            ["2", r"D:\论文111\data\WAMI代码讲解给导师版.docx", "你给导师讲代码时看的讲解稿。"],
            ["3", r"D:\论文111\data\method_audit_excels_expanded\WAMI_paper_faithful.xlsx", "WAMI 主方法逐样本证据。"],
            ["4", r"D:\论文111\data\method_audit_excels_expanded\WAMI_live_agent_action_level_qwen25.xlsx", "live-agent 动作级证据。"],
            ["5", r"D:\论文111\data\method_audit_excels_expanded", "导师要看 baseline 细节时再发整个文件夹。"],
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
